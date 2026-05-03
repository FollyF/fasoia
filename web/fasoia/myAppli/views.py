from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, authenticate, logout
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators import csrf
from django.views.decorators.csrf import csrf_protect
from django.views.decorators.http import require_http_methods
from django.conf import settings
from urllib.parse import quote
from django.utils import timezone
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import UploadedFile
from django.http import FileResponse
from datetime import datetime, timedelta

from docx import Document as DocxDocument
from io import BytesIO
from weasyprint import HTML
from html2docx import html2docx
from django.core.mail import EmailMessage
from .forms import *
from .models import *
from analyse_ia.models import *
from analyse_ia.recommandation_emploi import mettre_a_jour_scores_dossier
from .services.generateur_cv_public import GenerateurCVPublic
from .generateur import GenerateurDocument
from analyse_ia.ia_client import IAClient
from django.db.models import Count, Q

import re
import os
import json
import csv
import logging
import tempfile
import subprocess
import zipfile
import requests as req_http

logger = logging.getLogger(__name__)

def test(request):
    return render(request, 'myAppli/test.html')

def home(request):
    """
    Page d'accueil avec les dernières opportunités
    """
    # Récupérer les 5 dernières offres et AMI
    dernieres_offres = Offre_uemoa.objects.all().order_by('-date_scraping')[:5]
    derniers_amis = Ami_uemoa.objects.all().order_by('-date_scraping')[:5]
    derniers_emplois = OffreEmploi.objects.all().order_by('-date_scraping')[:5]
    
    # Compter le total
    total_opportunites = Offre_uemoa.objects.count() + Ami_uemoa.objects.count() + OffreEmploi.objects.count()
    
    context = {
        'dernieres_offres': dernieres_offres,
        'derniers_amis': derniers_amis,
        'derniers_emplois': derniers_emplois,
        'total_opportunites': total_opportunites,
    }
    return render(request, 'myAppli/home.html', context)

def opportunites(request):
    """
    Page publique - Affiche TOUTES les opportunités sans filtre
    """
    offres_uemoa = Offre_uemoa.objects.all().order_by('-date_scraping')
    amis_uemoa = Ami_uemoa.objects.all().order_by('-date_scraping')
    offres_emploi = OffreEmploi.objects.filter(statut='PUBLIEE', est_active=True).order_by('-date_publication')
    
    context = {
        'offres': offres_uemoa,
        'amis': amis_uemoa,
        'offres_emploi': offres_emploi,
        'total_opportunites': offres_uemoa.count() + amis_uemoa.count() + offres_emploi.count(),
    }
    return render(request, 'myAppli/opportunites.html', context)

@login_required
def dashboard_opportunites(request):
    """
    Page dashboard - Affiche les opportunités filtrées selon le profil utilisateur
    """
    offres_uemoa = Offre_uemoa.objects.none()
    amis_uemoa = Ami_uemoa.objects.none()
    offres_emploi = OffreEmploi.objects.none()
    
    # ===== PROFIL ENTREPRISE =====
    if hasattr(request.user, 'entreprise'):
        entreprise = request.user.entreprise
        offres_uemoa = Offre_uemoa.objects.all()
        amis_uemoa = Ami_uemoa.objects.all()
        
    # ===== PROFIL CANDIDAT =====
    elif hasattr(request.user, 'particulier') and hasattr(request.user.particulier, 'candidat'):
        candidat = request.user.particulier.candidat
        offres_emploi = OffreEmploi.objects.filter(statut='PUBLIEE', est_active=True)
        
    # ===== PROFIL RECRUTEUR =====
    elif hasattr(request.user, 'particulier') and hasattr(request.user.particulier, 'recruteur'):
        recruteur = request.user.particulier.recruteur
        offres_emploi = OffreEmploi.objects.filter(
            recruteur=recruteur,
            statut='PUBLIEE'
        ).order_by('-date_publication')
    
    # ===== PARTICULIER SANS RÔLE =====
    elif hasattr(request.user, 'particulier'):
        offres_emploi = OffreEmploi.objects.filter(statut='PUBLIEE', est_active=True).order_by('-date_publication')
    
    total_opportunites = offres_uemoa.count() + amis_uemoa.count() + offres_emploi.count()
    
    context = {
        'offre': offres_uemoa,
        'ami': amis_uemoa,
        'offres_emploi': offres_emploi,
        'total_opportunites': total_opportunites,
    }
    return render(request, 'myAppli/dashboard_opportunites.html', context)

@require_http_methods(["GET", "POST"])
@csrf_protect
def inscription(request):
    """
    Vue d'inscription - Gère les 3 types de profils
    """
    if request.user.is_authenticated:
    # Utilisateur déjà connecté - rediriger vers son dashboard
        if hasattr(request.user, 'entreprise'):
            return redirect('myAppli:dashboard_entreprise')
        elif hasattr(request.user, 'particulier'):
            particulier = request.user.particulier
            if hasattr(particulier, 'candidat') and hasattr(particulier, 'recruteur'):
                return redirect('myAppli:dashboard_particulier')
            elif hasattr(particulier, 'candidat'):
                return redirect('myAppli:dashboard_candidat')
            elif hasattr(particulier, 'recruteur'):
                return redirect('myAppli:dashboard_recruteur')
            else:
                return redirect('myAppli:dashboard_particulier')
        return redirect('myAppli:home')

    if request.method == 'POST':
        form = InscriptionForm(request.POST)
        
        if form.is_valid():
            try:
                user = form.save()
                
                # Connexion automatique après inscription
                login(request, user, backend='django.contrib.auth.backends.ModelBackend')
                
                # Message personnalisé selon le type de profil
                profile_type = form.cleaned_data['profile_type']
                
                if profile_type == 'entreprise':
                    entreprise = user.entreprise
                    messages.success(
                        request, 
                        f"Bienvenue {entreprise.raisonSociale} ! Votre compte entreprise a été créé. Complétez votre profil pour commencer."
                    )
                    logger.info(f"Nouvelle inscription entreprise : {user.email} - {entreprise.raisonSociale}")
                    return redirect('myAppli:dashboard_entreprise')
                    
                elif profile_type == 'particulier':
                    particulier = user.particulier
                    messages.success(
                        request, 
                        f"Bienvenue {particulier.prenom} {particulier.nom} ! Votre compte a été créé avec succès."
                    )
                    logger.info(f"Nouvelle inscription particulier : {user.email}")
                    return redirect('myAppli:dashboard_particulier')
                    
                elif profile_type == 'partenaire':
                    # Si partenaire = recruteur ou autre
                    partenaire = user.recruteur  # À adapter selon ton modèle
                    messages.success(
                        request, 
                        f"Bienvenue ! Votre compte partenaire a été créé avec succès."
                    )
                    logger.info(f"Nouvelle inscription partenaire : {user.email}")
                    return redirect('myAppli:home')
                
            except Exception as e:
                logger.error(f"Erreur lors de l'inscription : {str(e)}")
                messages.error(request, "Une erreur est survenue. Veuillez réessayer.")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, error)
    else:
        # Initialiser le formulaire avec le type par défaut
        initial_data = {'profile_type': 'particulier'}
        form = InscriptionForm(initial=initial_data)
    
    return render(request, 'myAppli/inscription.html', {
        'form': form,
        'title': 'Inscription',
        'no_header_footer': True
    })

@require_http_methods(["GET", "POST"])
@csrf_protect
def connexion(request):
    """
    Vue de connexion
    """
    print("="*50)
    print("VUE CONNEXION APPELEE")
    print(f"Utilisateur authentifié: {request.user.is_authenticated}")
    print(f"Utilisateur: {request.user}")
    print(f"Méthode: {request.method}")
    print("="*50)

    if request.user.is_authenticated:
    # Utilisateur déjà connecté - rediriger vers son dashboard
        if hasattr(request.user, 'entreprise'):
            return redirect('myAppli:dashboard_entreprise')
        elif hasattr(request.user, 'particulier'):
            particulier = request.user.particulier
            if hasattr(particulier, 'candidat') and hasattr(particulier, 'recruteur'):
                return redirect('myAppli:dashboard_particulier')
            elif hasattr(particulier, 'candidat'):
                return redirect('myAppli:dashboard_candidat')
            elif hasattr(particulier, 'recruteur'):
                return redirect('myAppli:dashboard_recruteur')
            else:
                return redirect('myAppli:dashboard_particulier')
        return redirect('myAppli:home')
 
    if request.method == 'POST':
        form = ConnexionForm(request, data=request.POST)
        
        if form.is_valid():
            user = form.get_user()
            
            # Gestion du "remember me"
            remember_me = form.cleaned_data.get('remember_me', False)
            if not remember_me:
                request.session.set_expiry(0)
            
            login(request, user)
            
            # Message personnalisé selon le type de profil
            if hasattr(user, 'entreprise'):
                messages.success(request, f"Bon retour parmi nous, {user.entreprise.raisonSociale} !")
            elif hasattr(user, 'particulier'):
                messages.success(request, f"Bon retour parmi nous, {user.particulier.prenom} !")
            elif hasattr(user, 'recruteur'):
                messages.success(request, f"Bon retour parmi nous, {user.recruteur.organisation} !")
            else:
                messages.success(request, f"Bon retour parmi nous, {user.email} !")
            
            logger.info(f"Connexion réussie : {user.email}")
            
            # Redirection selon le type de profil
            next_url = request.GET.get('next')
            if next_url:
                return redirect(next_url)
            elif hasattr(user, 'entreprise'):
                return redirect('myAppli:dashboard_entreprise')
            elif hasattr(user, 'particulier'):
                particulier = user.particulier
                if hasattr(particulier, 'candidat') and hasattr(particulier, 'recruteur'):
                    return redirect('myAppli:dashboard_particulier')
                elif hasattr(particulier, 'candidat'):
                    return redirect('myAppli:dashboard_candidat')
                elif hasattr(particulier, 'recruteur'):
                    return redirect('myAppli:dashboard_recruteur')
                else:
                    return redirect('myAppli:dashboard_particulier')
            else:
                return redirect('myAppli:home')
        else:
            logger.warning(f"Tentative de connexion échouée pour : {request.POST.get('username', 'inconnu')}")
            messages.error(request, "Email ou mot de passe incorrect")
    else:
        form = ConnexionForm()
    
    return render(request, 'myAppli/connexion.html', {
        'form': form,
        'title': 'Connexion',
        'no_header_footer': True
    })

@login_required
def deconnexion(request):
    """
    Vue de déconnexion
    """
    email = request.user.email
    logout(request)
    messages.success(request, "Vous avez été déconnecté avec succès")
    logger.info(f"Déconnexion : {email}")
    return redirect('myAppli:connexion')

@login_required
def profil(request):
    """
    Vue du profil utilisateur (générique)
    """
    context = {'user': request.user}
    
    # Ajouter les données spécifiques selon le type
    if hasattr(request.user, 'entreprise'):
        context['entreprise'] = request.user.entreprise
        context['profil_type'] = 'entreprise'
    elif hasattr(request.user, 'particulier'):
        context['particulier'] = request.user.particulier
        context['profil_type'] = 'particulier'
    elif hasattr(request.user, 'recruteur'):
        context['recruteur'] = request.user.recruteur
        context['profil_type'] = 'recruteur'
    
    return render(request, 'myAppli/profil.html', context)

@login_required
def dashboard_entreprise(request):
    """
    Tableau de bord spécifique pour les entreprises
    """
    # Vérifier que l'utilisateur est bien une entreprise
    if not hasattr(request.user, 'entreprise'):
        messages.error(request, "Accès réservé aux entreprises")
        return redirect('myAppli:home')
    
    entreprise = request.user.entreprise
    
    # ===== CODE POUR LE CALCUL DE PROGRESSION =====
    champs_obligatoires = [
        'domaineActive', 'localisation', 'competencesCles',
        'pays_intervention', 'chiffre_affaires', 'types_opportunites'
    ]
    
    champs_remplis = 0
    details_champs = {}
    
    for champ in champs_obligatoires:
        valeur = getattr(entreprise, champ)
        
        if champ == 'types_opportunites' or champ == 'pays_intervention':
            est_rempli = bool(valeur)
        elif champ == 'chiffre_affaires':
            est_rempli = valeur is not None and valeur > 0
        else:
            est_rempli = bool(valeur and str(valeur).strip())
        
        if est_rempli:
            champs_remplis += 1
            details_champs[champ] = True
        else:
            details_champs[champ] = False
    
    total_champs = len(champs_obligatoires)
    pourcentage = int((champs_remplis / total_champs) * 100) if total_champs > 0 else 0
    champs_manquants = total_champs - champs_remplis
    profil_complet = (champs_manquants == 0)
    
    champs_manquants_liste = []
    libelles_champs = {
        'domaineActive': "Domaine d'activité",
        'localisation': "Localisation",
        'competencesCles': "Compétences clés",
        'pays_intervention': "Pays d'intervention",
        'chiffre_affaires': "Chiffre d'affaires",
        'types_opportunites': "Types d'opportunités"
    }
    
    for champ in champs_obligatoires:
        if not details_champs[champ]:
            champs_manquants_liste.append(libelles_champs.get(champ, champ))
    
    # ===== CODE POUR LES RECOMMANDATIONS (CORRIGÉ) =====
    recommandations = []
    opportunites_correspondantes = 0
    recommandations_count = 0
    
    if profil_complet:
        try:
            from analyse_ia.models import Recommandation  # ← Correction ici !
            
            recommandations = Recommandation.objects.filter(
                entreprise=entreprise
            ).order_by('-score_global')[:10]

            for reco in recommandations:
                reco.score_global = round(reco.score_global * 100, 1)
            
            recommandations_count = recommandations.count()
            opportunites_correspondantes = recommandations_count
            
        except Exception as e:
            print(f"Erreur lors du chargement des recommandations: {e}")
            recommandations = []
            recommandations_count = 0
            opportunites_correspondantes = 0
    
    return render(request, 'myAppli/dashboard_entreprise.html', {
        'entreprise': entreprise,
        'user': request.user,
        'pourcentage_completion': pourcentage,
        'champs_remplis': champs_remplis,
        'total_champs': total_champs,
        'champs_obligatoires_manquants': champs_manquants,
        'profil_complet': profil_complet,
        'champs_manquants_liste': champs_manquants_liste,
        'details_champs': details_champs,
        'recommandations': recommandations,
        'opportunites_correspondantes': opportunites_correspondantes,
        'recommandations_count': recommandations_count,
    })

@login_required
def dashboard_particulier(request):
    """
    Dashboard principal pour les particuliers
    Permet de choisir son rôle (candidat ou recruteur)
    """
    # Vérifier que l'utilisateur est bien un particulier
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Accès réservé aux particuliers")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    # Vérifier si l'utilisateur a déjà des profils
    a_profil_candidat = hasattr(particulier, 'candidat')
    a_profil_recruteur = hasattr(particulier, 'recruteur')
    
    context = {
        'particulier': particulier,
        'a_profil_candidat': a_profil_candidat,
        'a_profil_recruteur': a_profil_recruteur,
    }
    
    return render(request, 'myAppli/dashboard_particulier.html', context)

@login_required
def dashboard_candidat(request):
    """
    Dashboard spécifique pour les demandeurs d'emploi
    """
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Accès non autorisé")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    if not hasattr(particulier, 'candidat'):
        messages.warning(request, "Vous devez d'abord activer votre profil candidat")
        return redirect('myAppli:dashboard_particulier')
    
    candidat = particulier.candidat
    
    # Récupérer les offres recommandées
    offres_recommandees = OffreEmploi.objects.filter(
        statut='PUBLIEE',
        est_active=True
    ).order_by('-date_publication')[:10]
    
    # ===== CALCUL CORRIGÉ DE LA PROGRESSION =====
    # Liste de TOUS les champs à prendre en compte (particulier + candidat)
    champs_profil = [
        # Champs du Particulier
        ('nom', particulier.nom),
        ('prenom', particulier.prenom),
        ('email', particulier.email),
        ('telephone', particulier.telephone),
        ('date_naissance', particulier.date_naissance),
        ('adresse', particulier.adresse),
        ('ville', particulier.ville),
        ('pays', particulier.pays),
        
        # Champs du Candidat
        ('niveauEtude', candidat.niveauEtude),
        ('competences', candidat.competences),
        ('disponibilite', candidat.disponibilite),
        ('niveauLangues', candidat.niveauLangues),
        ('secteur_recherche', candidat.secteur_recherche),
        ('type_contrat_recherche', candidat.type_contrat_recherche),
        ('localisation_recherche', candidat.localisation_recherche),
        ('anneesExperiences', candidat.anneesExperiences),
        ('salaire_souhaite', candidat.salaire_souhaite),
        ('mobilite', candidat.mobilite),  # True/False donc toujours rempli
        ('cv', candidat.cv),
        ('lettre_motivation', candidat.lettre_motivation),
    ]
    
    # Compter les champs remplis
    champs_remplis = 0
    details_champs = {}
    
    for nom_champ, valeur in champs_profil:
        # Déterminer si le champ est considéré comme rempli
        if nom_champ == 'mobilite':
            # La mobilité est un boolean, toujours rempli (True ou False)
            est_rempli = True
        elif nom_champ in ['cv', 'lettre_motivation']:
            # Les fichiers sont considérés comme remplis si ils existent
            est_rempli = bool(valeur)
        elif nom_champ in ['anneesExperiences', 'salaire_souhaite']:
            # Les nombres : >0 est considéré comme rempli
            est_rempli = valeur is not None
        elif nom_champ == 'date_naissance':
            # Date de naissance optionnelle
            est_rempli = bool(valeur)
        else:
            # Champs texte : non vides et non "None"
            est_rempli = bool(valeur and str(valeur).strip())
        
        if est_rempli:
            champs_remplis += 1
            details_champs[nom_champ] = True
        else:
            details_champs[nom_champ] = False
    
    total_champs = len(champs_profil)
    progression = int((champs_remplis / total_champs) * 100) if total_champs > 0 else 0
    
    # Définir les champs obligatoires (ceux avec * dans le formulaire)
    champs_obligatoires = [
        ('nom', particulier.nom),
        ('prenom', particulier.prenom),
        ('email', particulier.email),
        ('telephone', particulier.telephone),
        ('niveauEtude', candidat.niveauEtude),
        ('competences', candidat.competences),
        ('disponibilite', candidat.disponibilite),
        ('cv', candidat.cv),
    ]
    
    # Vérifier si tous les champs obligatoires sont remplis
    profil_complet = True
    champs_manquants = []
    
    for nom_champ, valeur in champs_obligatoires:
        if nom_champ == 'cv':
            # CV uploadé OU CV généré avec PDF
            a_un_cv = bool(candidat.cv) or CVGenere.objects.filter(
                utilisateur=candidat.particulier.user,
                fichier_pdf__isnull=False
            ).exists()
            if not a_un_cv:
                profil_complet = False
                champs_manquants.append('cv')
        else:
            if not bool(valeur and str(valeur).strip()):
                profil_complet = False
                champs_manquants.append(nom_champ)
    
    # Gérer l'affichage du formulaire
    show_form = not profil_complet or request.GET.get('edit') == 'true'
    
    # Si on vient juste d'enregistrer, on cache le formulaire
    if request.session.pop('profil_juste_complete', False):
        show_form = False
    
    # Libellés des champs pour l'affichage
    libelles_champs = {
        'nom': 'Nom',
        'prenom': 'Prénom',
        'email': 'Email',
        'telephone': 'Téléphone',
        'date_naissance': 'Date de naissance',
        'adresse': 'Adresse',
        'ville': 'Ville',
        'pays': 'Pays',
        'niveauEtude': "Niveau d'étude",
        'competences': 'Compétences',
        'disponibilite': 'Disponibilité',
        'niveauLangues': 'Niveau en langues',
        'secteur_recherche': 'Secteur recherché',
        'type_contrat_recherche': 'Type de contrat',
        'localisation_recherche': 'Localisation recherchée',
        'anneesExperiences': "Années d'expérience",
        'salaire_souhaite': 'Salaire souhaité',
        'mobilite': 'Mobilité',
        'cv': 'CV',
        'lettre_motivation': 'Lettre de motivation',
    }
    
    offres_recommandees = []

    if profil_complet:
        try:
            recommandations = RecommandationEmploi.objects.filter(
                candidat=candidat
            ).select_related('offre').order_by('-score_global')[:10]

            for reco in recommandations:
                offre = reco.offre
                offre.score = round(reco.score_global * 100, 1)
                offres_recommandees.append(offre)

        except Exception as e:
            print(f"Erreur chargement recommandations: {e}")
            offres_recommandees = list(OffreEmploi.objects.filter(
                statut='PUBLIEE',
                est_active=True
            ).order_by('-date_publication')[:10])
    else:
        offres_recommandees = list(OffreEmploi.objects.filter(
            statut='PUBLIEE',
            est_active=True
        ).order_by('-date_publication')[:5])

    cvs_generes = CVGenere.objects.filter(
        utilisateur=candidat.particulier.user,
        fichier_pdf__isnull=False
    ).order_by('-date_generation')

    cv_actif = (
        cvs_generes.filter(est_utilise=True).first() or
        cvs_generes.filter(est_favori=True).first() or
        cvs_generes.first()
    )

    context = {
        'candidat': candidat,
        'particulier': particulier,
        'offres_recommandees': offres_recommandees,
        'progression': progression,
        'champs_remplis': champs_remplis,
        'total_champs': total_champs,
        'profil_complet': profil_complet,
        'champs_manquants': champs_manquants,
        'libelles_champs': libelles_champs,
        'details_champs': details_champs,
        'show_form': show_form,
        'recommandations_count': len(offres_recommandees),
        'cvs_generes': cvs_generes,  
        'cv_actif': cv_actif,  
    }
    
    return render(request, 'myAppli/dashboard_candidat.html', context)

@login_required
def dashboard_recruteur(request):
    """
    Dashboard spécifique pour les recruteurs
    """
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Accès non autorisé")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    if not hasattr(particulier, 'recruteur'):
        messages.warning(request, "Vous devez d'abord activer votre profil recruteur")
        return redirect('myAppli:dashboard_particulier')
    
    recruteur = particulier.recruteur
    
    # Offres du recruteur
    offres_publiees = OffreEmploi.objects.filter(
        recruteur=recruteur, 
        statut='PUBLIEE'
    ).annotate(
        nb_candidats_reels=Count(
            'dossiers_candidature',
            filter=Q(dossiers_candidature__statut__in=['SOUMIS', 'VU'])
        )
    ).order_by('-date_publication')
    
    # ✅ RECOMMANDATIONS = candidats qui ont postulé, triés par score IA
    talents_recommandes = DossierCandidature.objects.filter(
        offre__recruteur=recruteur,
        statut__in=[DossierCandidature.SOUMIS, DossierCandidature.VU]
    ).select_related('offre', 'candidat__particulier').order_by('offre__titre', '-score_compatibilite')[:20]
    
    # Statistiques
    stats = {
        'offres_publiees': offres_publiees.count(),
        'candidatures_recues': DossierCandidature.objects.filter(
            offre__recruteur=recruteur,
            statut__in=[DossierCandidature.SOUMIS, DossierCandidature.VU]
        ).count(),
        'talents_recommandes': talents_recommandes.count(),
    }
    
    # Progression du profil
    champs_obligatoires = ['organisation', 'secteur', 'typeStructure', 'poste_occupe']
    champs_remplis = sum(1 for champ in champs_obligatoires 
                        if getattr(recruteur, champ) and str(getattr(recruteur, champ)).strip())
    
    total_champs = len(champs_obligatoires)
    progression = int((champs_remplis / total_champs) * 100) if total_champs > 0 else 0
    profil_complet = (champs_remplis == total_champs)
    
    context = {
        'recruteur': recruteur,
        'offres_publiees': offres_publiees,
        'talents_recommandes': talents_recommandes,
        'stats': stats,
        'progression': progression,
        'champs_remplis': champs_remplis,
        'total_champs': total_champs,
        'profil_complet': profil_complet,
    }
    
    return render(request, 'myAppli/dashboard_recruteur.html', context)
    
@login_required
def activer_profil_candidat(request):
    """Active le profil candidat pour un particulier"""
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Vous devez être un particulier")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    # Vérifier si le profil existe déjà
    if hasattr(particulier, 'candidat'):
        messages.info(request, "Vous avez déjà un profil candidat")
        return redirect('myAppli:dashboard_candidat')
    
    # Créer le profil candidat
    Candidat.objects.create(
        particulier=particulier,
        niveauEtude='',
        competences='',
        disponibilite='',
        secteur_recherche='',
        type_contrat_recherche='',
        localisation_recherche='',
        mobilite=False
    )
    
    messages.success(request, "Profil candidat activé avec succès !")
    return redirect('myAppli:dashboard_candidat')

@login_required
def activer_profil_recruteur(request):
    """Active le profil recruteur pour un particulier"""
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Vous devez être un particulier")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    # Vérifier si le profil existe déjà
    if hasattr(particulier, 'recruteur'):
        messages.info(request, "Vous avez déjà un profil recruteur")
        return redirect('myAppli:dashboard_recruteur')
    
    # Créer le profil recruteur
    Recruteur.objects.create(
        particulier=particulier,
        organisation='',
        secteur='',
        typeStructure='',
        poste_occupe=''
    )
    
    messages.success(request, "Profil recruteur activé avec succès !")
    return redirect('myAppli:dashboard_recruteur')

@login_required
@require_http_methods(["POST"])
@csrf_protect
def completer_profil_candidat(request):
    """
    Vue pour enregistrer les modifications du profil candidat
    """
    print("="*50)
    print("DONNÉES REÇUES PROFIL CANDIDAT:")
    print(request.POST)
    print(request.FILES)
    print("="*50)
    
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Accès réservé aux particuliers")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    if not hasattr(particulier, 'candidat'):
        messages.warning(request, "Vous devez d'abord activer votre profil candidat")
        return redirect('myAppli:dashboard_particulier')
    
    candidat = particulier.candidat
    
    try:
        # ===== MISE À JOUR DES CHAMPS DU PARTICULIER =====
        particulier.nom = request.POST.get('nom', particulier.nom)
        particulier.prenom = request.POST.get('prenom', particulier.prenom)
        particulier.email = request.POST.get('email', particulier.email)
        particulier.telephone = request.POST.get('telephone', particulier.telephone)
        
        # Date de naissance (optionnelle)
        date_naissance = request.POST.get('date_naissance')
        if date_naissance:
            try:
                from datetime import datetime
                particulier.date_naissance = datetime.strptime(date_naissance, '%Y-%m-%d').date()
            except ValueError:
                particulier.date_naissance = None
                
        particulier.adresse = request.POST.get('adresse', particulier.adresse)
        particulier.ville = request.POST.get('ville', particulier.ville)
        particulier.pays = request.POST.get('pays', particulier.pays)
        
        # Sauvegarder le particulier
        particulier.save()
        print("✅ Particulier mis à jour")
        
        # ===== MISE À JOUR DES CHAMPS DU CANDIDAT =====
        # Gestion des chaînes vides pour les champs texte
        candidat.niveauEtude = request.POST.get('niveauEtude', '')
        candidat.competences = request.POST.get('competences', '')
        candidat.disponibilite = request.POST.get('disponibilite', '')
        candidat.niveauLangues = request.POST.get('niveauLangues', '')
        candidat.secteur_recherche = request.POST.get('secteur_recherche', '')
        candidat.type_contrat_recherche = request.POST.get('type_contrat_recherche', '')
        candidat.localisation_recherche = request.POST.get('localisation_recherche', '')
        
        # ✅ Gestion des nombres avec valeur par défaut 0
        annees = request.POST.get('anneesExperiences', '0')
        candidat.anneesExperiences = int(annees) if annees and annees.strip() else 0
        
        # ✅ Gestion du salaire (optionnel)
        salaire = request.POST.get('salaire_souhaite', '')
        if salaire and salaire.strip():
            try:
                candidat.salaire_souhaite = float(salaire)
            except ValueError:
                candidat.salaire_souhaite = None
        else:
            candidat.salaire_souhaite = None
        
        # ✅ Gestion de la mobilité (checkbox)
        candidat.mobilite = request.POST.get('mobilite') == 'on'
        
        # ✅ Gestion des fichiers
        if 'cv' in request.FILES:
            candidat.cv = request.FILES['cv']
        if 'lettre_motivation' in request.FILES:
            candidat.lettre_motivation = request.FILES['lettre_motivation']
        
        candidat.save()
        print("✅ Candidat mis à jour")
        
        # ===== VÉRIFICATION PROFIL COMPLET =====
        champs_obligatoires = {
            'niveauEtude': candidat.niveauEtude,
            'competences': candidat.competences,
            'anneesExperiences': candidat.anneesExperiences,
            'secteur_recherche': candidat.secteur_recherche,
            'localisation_recherche': candidat.localisation_recherche,
            'niveauLangues': candidat.niveauLangues,
        }

        profil_complet = True
        champs_manquants = []

        for champ, valeur in champs_obligatoires.items():
            if not valeur or str(valeur).strip() == '' or valeur == 0:
                profil_complet = False
                champs_manquants.append(champ)

        # Vérifier le CV — uploadé OU généré
        from myAppli.models import CVGenere
        a_un_cv = (
            candidat.cv or  # CV uploadé manuellement
            CVGenere.objects.filter(
                utilisateur=candidat.particulier.user,
                fichier_pdf__isnull=False
            ).exists()  # CV généré avec PDF
        )

        if not a_un_cv:
            profil_complet = False
            champs_manquants.append('cv')

        print(f"Profil complet: {profil_complet}")
        if champs_manquants:
            print(f"Champs manquants: {champs_manquants}")

        if profil_complet:
            messages.success(request, "Félicitations ! Votre profil est complet. Vos recommandations sont en cours de génération.")
            try:
                from analyse_ia.moteur_recommandation_emploi import MoteurRecommandationEmploi
                moteur = MoteurRecommandationEmploi()
                moteur.recommander_pour_candidat(candidat.particulier_id)
                messages.info(request, "Vos recommandations d'offres d'emploi ont été générées !")
            except Exception as e:
                print(f"Erreur recommandations (non bloquante): {e}")
        else:
            messages.success(request, "Profil mis à jour ! Complétez tous les champs pour activer les recommandations.")

    except Exception as e:
        print(f"❌ ERREUR: {str(e)}")
        import traceback
        traceback.print_exc()
        messages.error(request, f"Erreur lors de la sauvegarde : {str(e)}")
    
    return redirect('myAppli:dashboard_candidat')

@login_required
@require_http_methods(["POST"])
@csrf_protect
def completer_profil_recruteur(request):
    """
    Vue pour enregistrer les modifications du profil recruteur
    """
    print("="*50)
    print("DONNÉES REÇUES PROFIL RECRUTEUR:")
    print(request.POST)
    print("="*50)
    
    if not hasattr(request.user, 'particulier'):
        messages.error(request, "Accès réservé aux particuliers")
        return redirect('myAppli:home')
    
    particulier = request.user.particulier
    
    if not hasattr(particulier, 'recruteur'):
        messages.warning(request, "Vous devez d'abord activer votre profil recruteur")
        return redirect('myAppli:dashboard_particulier')
    
    recruteur = particulier.recruteur
    
    try:
        # Mise à jour des champs simples
        recruteur.organisation = request.POST.get('organisation', recruteur.organisation)
        recruteur.secteur = request.POST.get('secteur', recruteur.secteur)
        recruteur.typeStructure = request.POST.get('typeStructure', recruteur.typeStructure)
        recruteur.poste_occupe = request.POST.get('poste_occupe', recruteur.poste_occupe)
        
        # Traitement des secteurs recherchés
        secteurs_recherches = request.POST.get('secteurs_recherches', '')
        if secteurs_recherches and secteurs_recherches.strip():
            recruteur.secteurs_recherches = [s.strip() for s in secteurs_recherches.split(',') if s.strip()]
        else:
            recruteur.secteurs_recherches = []
        
        # Traitement des types de contrats
        types_contrats = request.POST.get('types_contrats_proposes', '')
        if types_contrats and types_contrats.strip():
            recruteur.types_contrats_proposes = [t.strip() for t in types_contrats.split(',') if t.strip()]
        else:
            recruteur.types_contrats_proposes = []
        
        recruteur.save()
        
        # ===== VÉRIFICATION DU PROFIL COMPLET =====
        champs_obligatoires = ['organisation', 'secteur', 'typeStructure', 'poste_occupe']
        champs_remplis = 0
        champs_manquants = []
        
        for champ in champs_obligatoires:
            valeur = getattr(recruteur, champ)
            if valeur and str(valeur).strip():
                champs_remplis += 1
            else:
                champs_manquants.append(champ)
        
        total_champs = len(champs_obligatoires)
        progression = int((champs_remplis / total_champs) * 100) if total_champs > 0 else 0
        profil_complet = (champs_remplis == total_champs)
        
        print(f"📊 Progression: {progression}% ({champs_remplis}/{total_champs})")
        print(f"✅ Profil complet: {profil_complet}")
        
        # Message personnalisé
        if profil_complet:
            messages.success(request, "🎉 Félicitations ! Votre profil recruteur est maintenant complet ! Vous pouvez maintenant publier des offres.")
        else:
            messages.success(request, f"Votre profil a été mis à jour ! Il vous manque : {', '.join(champs_manquants)}")
        
    except Exception as e:
        print(f"❌ ERREUR: {str(e)}")
        import traceback
        traceback.print_exc()
        messages.error(request, f"Erreur : {str(e)}")
    
    return redirect('myAppli:dashboard_recruteur')

@login_required
@require_http_methods(["POST"])
@csrf_protect
def completer_profil_entreprise(request):
    """
    Vue pour enregistrer les modifications du profil entreprise
    """
    print("="*50)
    print("DONNÉES REÇUES:")
    print(request.POST)
    print("="*50)
    
    if not hasattr(request.user, 'entreprise'):
        messages.error(request, "Accès réservé aux entreprises")
        return redirect('myAppli:home')
    
    entreprise = request.user.entreprise
    
    try:
        # Mise à jour des champs simples
        entreprise.domaineActive = request.POST.get('domaineActive', '')
        entreprise.localisation = request.POST.get('localisation', '')
        entreprise.competencesCles = request.POST.get('competencesCles', '')
        
        # Gestion des nombres - NE PAS METTRE DE CHAÎNE VIDE
        taille = request.POST.get('taille')
        entreprise.taille = int(taille) if taille and taille.strip() else 0
        
        annee = request.POST.get('annee_creation')
        entreprise.annee_creation = int(annee) if annee and annee.strip() else None
        
        site_web = request.POST.get('site_web')
        entreprise.site_web = site_web if site_web else ''
        
        description = request.POST.get('description')
        entreprise.description = description if description else ''
        
        annees_exp = request.POST.get('annees_experience')
        entreprise.annees_experience = int(annees_exp) if annees_exp and annees_exp.strip() else 0
        
        nb_projets = request.POST.get('nb_projets_realises')
        entreprise.nb_projets_realises = int(nb_projets) if nb_projets and nb_projets.strip() else 0
        
        references = request.POST.get('references')
        entreprise.references = references if references else ''
        
        rayon = request.POST.get('rayon_action')
        entreprise.rayon_action = int(rayon) if rayon and rayon.strip() else None
        
        montant_min = request.POST.get('montant_min')
        entreprise.montant_min = float(montant_min) if montant_min and montant_min.strip() else None
        
        montant_max = request.POST.get('montant_max')
        entreprise.montant_max = float(montant_max) if montant_max and montant_max.strip() else None
        
        # Chiffre d'affaires (convertir en Decimal)
        ca = request.POST.get('chiffre_affaires')
        if ca and ca.strip():
            entreprise.chiffre_affaires = ca
        
        # Capital social
        cs = request.POST.get('capital_social')
        if cs and cs.strip():
            entreprise.capital_social = cs
        
        # Traitement des champs JSON (listes)
        # Pays d'intervention
        pays = request.POST.get('pays_intervention', '')
        if pays and pays.strip():
            entreprise.pays_intervention = [p.strip() for p in pays.split(',') if p.strip()]
        else:
            entreprise.pays_intervention = []
        
        # Certifications
        certifs = request.POST.get('certifications', '')
        if certifs and certifs.strip():
            entreprise.certifications = [c.strip() for c in certifs.split(',') if c.strip()]
        else:
            entreprise.certifications = []
        
        # Agréments
        agrements = request.POST.get('agrements', '')
        if agrements and agrements.strip():
            entreprise.agrements = [a.strip() for a in agrements.split(',') if a.strip()]
        else:
            entreprise.agrements = []
        
        # Types d'opportunités (select multiple)
        types_opportunites = request.POST.getlist('types_opportunites')
        if types_opportunites:
            entreprise.types_opportunites = list(types_opportunites)
        else:
            entreprise.types_opportunites = []
        
        print(f"AVANT SAUVEGARDE: domaineActive={entreprise.domaineActive}")
        print(f"AVANT SAUVEGARDE: annee_creation={entreprise.annee_creation}")
        
        # Sauvegarder
        entreprise.save()
        print("APRÈS SAUVEGARDE: OK")
        
        # Mettre à jour l'index des mots-clés
        try:
            entreprise.sauvegarder_mots_cles()
            print("Mots-clés sauvegardés")
        except Exception as e:
            print(f"Erreur mots-clés (non bloquante): {e}")
        
        # ===== NOUVELLE VÉRIFICATION CORRIGÉE =====
        # Vérifier si le profil est maintenant complet
        profil_complet = True
        champs_manquants = []
        
        # 1. Domaine d'activité
        if not entreprise.domaineActive or not str(entreprise.domaineActive).strip():
            profil_complet = False
            champs_manquants.append("Domaine d'activité")
        
        # 2. Localisation
        if not entreprise.localisation or not str(entreprise.localisation).strip():
            profil_complet = False
            champs_manquants.append("Localisation")
        
        # 3. Compétences clés
        if not entreprise.competencesCles or not str(entreprise.competencesCles).strip():
            profil_complet = False
            champs_manquants.append("Compétences clés")
        
        # 4. Pays d'intervention (liste)
        if not entreprise.pays_intervention or len(entreprise.pays_intervention) == 0:
            profil_complet = False
            champs_manquants.append("Pays d'intervention")
        
        # 5. Chiffre d'affaires (nombre)
        try:
            ca_value = float(entreprise.chiffre_affaires) if entreprise.chiffre_affaires else 0
            if ca_value <= 0:
                profil_complet = False
                champs_manquants.append("Chiffre d'affaires")
        except (ValueError, TypeError):
            profil_complet = False
            champs_manquants.append("Chiffre d'affaires (invalide)")
        
        # 6. Types d'opportunités (liste)
        if not entreprise.types_opportunites or len(entreprise.types_opportunites) == 0:
            profil_complet = False
            champs_manquants.append("Types d'opportunités")
        
        print(f"Profil complet: {profil_complet}")
        if not profil_complet:
            print(f"Champs manquants: {', '.join(champs_manquants)}")
        # ===== FIN DE LA VÉRIFICATION CORRIGÉE =====
        
        if profil_complet:
            messages.success(request, "Félicitations ! Votre profil est maintenant complet. Vous allez recevoir des recommandations personnalisées.")
            
            # Optionnel : Lancer immédiatement les recommandations
            try:
                from analyse_ia.moteur_recommandation import MoteurRecommandation
                moteur = MoteurRecommandation()
                moteur.recommander_pour_entreprise(entreprise)
                messages.info(request, "Vos recommandations ont été générées.")
            except Exception as e:
                print(f"Erreur lors de la génération des recommandations: {e}")
        else:
            messages.success(request, "Profil mis à jour avec succès ! Continuez à le compléter pour activer les recommandations.")
        
    except Exception as e:
        print(f"ERREUR: {str(e)}")
        messages.error(request, f"Erreur lors de la sauvegarde : {str(e)}")
    
    return redirect('myAppli:dashboard_entreprise')

def detail_offre(request, pk):
    """
    Vue pour afficher les détails d'une offre
    """
    try:
        offre = Offre_uemoa.objects.get(pk=pk)
        return render(request, 'myAppli/detail_offre.html', {'offre': offre})
    except Offre_uemoa.DoesNotExist:
        messages.error(request, "Cette offre n'existe pas.")
        return redirect('myAppli:opportunites')

def detail_ami(request, pk):
    """
    Vue pour afficher les détails d'un AMI
    """
    try:
        ami = Ami_uemoa.objects.get(pk=pk)
        return render(request, 'myAppli/detail_ami.html', {'ami': ami})
    except Ami_uemoa.DoesNotExist:
        messages.error(request, "Cet AMI n'existe pas.")
        return redirect('myAppli:opportunites')

def detail_emploi(request, pk):
    """
    Vue pour afficher les détails d'une offre d'emploi
    """
    try:
        emploi = OffreEmploi.objects.get(pk=pk)
        return render(request, 'myAppli/detail_ami.html', {'emploi': emploi})
    except OffreEmploi.DoesNotExist:
        messages.error(request, "Cette offre d'emploi n'existe pas.")
        return redirect('myAppli:opportunites')

@login_required
def modifier_offre(request, offre_id):
    # 1. On récupère l'offre ou on renvoie une erreur 404
    offre = get_object_or_404(
        OffreEmploi, 
        id=offre_id, 
        recruteur__particulier__user=request.user
    )
    
    if request.method == 'POST':
        # 2. On lie le formulaire à l'instance existante avec les nouvelles données (POST)
        # On ajoute request.FILES si tu as des fichiers (ex: fichier_local)
        form = OffreEmploiForm(request.POST, request.FILES, instance=offre)
        
        if form.is_valid():
            form.save()
            messages.success(request, "L'offre a été mise à jour avec succès !")
            return redirect('myAppli:dashboard_recruteur') # Redirige vers ton tableau de bord
    else:
        # 3. Mode lecture : on affiche le formulaire pré-rempli
        form = OffreEmploiForm(instance=offre)
    
    return render(request, 'myAppli/outils_emploi/modifier_offre.html', {
        'form': form,
        'offre': offre
    })

@login_required
def nouvelle_soumission(request, opportunite_type, opportunite_id):
    """
    Page principale de génération de dossier de candidature
    Version propre - chaque offre a son propre dossier
    
    Pour l'instant : uniquement les AMI (Appels à Manifestation d'Intérêt)
    Les appels d'offres seront ajoutés ultérieurement
    """
    
    # ============================================
    # 1. Vérifier que c'est un AMI (pour l'instant)
    # ============================================
    if opportunite_type != 'Ami_uemoa':
        messages.warning(
            request, 
            "La soumission pour les appels d'offres sera disponible prochainement. "
            "Pour l'instant, seuls les AMI sont supportés."
        )
        return redirect('myAppli:opportunites')
    
    # ============================================
    # 2. Récupérer l'opportunité (AMI)
    # ============================================
    opportunite = get_object_or_404(Ami_uemoa, id=opportunite_id)
    type_nom = "Appel à manifestation d'intérêt"
    
    # ============================================
    # 3. Récupérer l'entreprise connectée
    # ============================================
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        messages.error(request, "Vous devez avoir une entreprise associée à votre compte.")
        return redirect('myAppli:home')
    
    # ============================================
    # 4. Récupérer ou créer le dossier de soumission
    # ============================================
    dossier, created = DossierSoumission.objects.get_or_create(
        entreprise=entreprise,
        opportunite_type=opportunite_type,
        opportunite_id=opportunite_id,
        defaults={
            'reference': f"AMI-{opportunite_id}-{entreprise.id}-{timezone.now().strftime('%Y%m%d%H%M%S')}",
            'statut': 'EN_PREPARATION',
            'date_soumission_prevue': timezone.now().date() + timezone.timedelta(days=30),
        }
    )
    
    # ============================================
    # 5. Définir les 6 documents requis pour un AMI
    # ============================================
    types_documents = [
        'ENVELOPPE',      # Message à coller sur l'enveloppe
        'LETTRE',         # Lettre de manifestation d'intérêt
        'PRESENTATION',   # Présentation de l'entreprise
        'FICHE',          # Fiche de renseignement
        'MATERIEL',       # Liste du matériel
        'PERSONNEL',      # Liste du personnel cadre
    ]
    
    # ============================================
    # 6. Créer automatiquement les documents manquants
    # ============================================
    if created:
        for type_doc in types_documents:
            DocumentGenere.objects.create(
                dossier=dossier,
                type_document=type_doc,
                statut='MISSING'
            )
    
    # ============================================
    # 7. Contexte pour le template
    # ============================================
    context = {
        'opportunite': opportunite,
        'opportunite_type': opportunite_type,
        'type_nom': type_nom,
        'dossier': dossier,
        'entreprise': entreprise,
        'types_documents': types_documents,
    }
    
    return render(request, 'myAppli/soumission/nouvelle_soumission.html', context)

@login_required
def mes_soumissions(request):
    """Affiche la liste de tous les dossiers de l'entreprise"""
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        messages.error(request, "Vous devez être une entreprise")
        return redirect('myAppli:home')
    
    # Récupérer TOUS les dossiers (sans filtre de statut)
    dossiers = DossierSoumission.objects.filter(
        entreprise=entreprise
    ).order_by('-date_creation')
    
    return render(request, 'myAppli/soumission/mes_soumissions.html', {'dossiers': dossiers})

def get_entete_image_html(request, entreprise):
    """Retourne le HTML de l'image d'en-tête avec URL absolue"""
    if entreprise.entete_image and entreprise.entete_image.url:
        image_url = request.build_absolute_uri(entreprise.entete_image.url)
        return f'<img src="{image_url}" style="width: 100%; margin-bottom: 20px;">'
    return ""

def construire_html_enveloppe(texte):
    """Construit le HTML pour le message enveloppe avec la mise en forme centrale"""
    texte_maj = texte.upper()
    return f"""
        <div style="font-family: Arial, sans-serif; padding: 20px; text-align: center;">
            <div style="font-weight: bold; font-size: 32px; line-height: 1.8;">
                {texte_maj}
            </div>
        </div>
        """

def construire_html_lettre(request, entreprise, description, responsable, date_actuelle, domaines):
    """Construit le HTML pour la lettre de manifestation"""
    
    entete_image_html = get_entete_image_html(request, entreprise)

    return f"""<div style="font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto;">
        {entete_image_html}
        <table style="width: 100%; border-collapse: collapse; margin-bottom: 30px;">
            <tr><td style="width: 50%; padding: 4px 0;"><strong>{entreprise.raisonSociale}</strong></td>
                <td style="width: 50%; padding: 4px 0; text-align: right;"><strong>BURKINA FASO</strong></td>
            </tr>
            <tr><td style="padding: 4px 0;">Tel: {entreprise.telephone or '+226 XX XX XX XX'}</td>
                <td style="padding: 4px 0; text-align: right;">La Patrie ou, La Mort, Nous vaincrons</td>
            </tr>
            <tr><td style="padding: 4px 0;">Email: {entreprise.email or 'contact@entreprise.com'}</td>
                <td style="padding: 4px 0; text-align: right;">A</td>
            </tr>
            <tr><td style="padding: 4px 0;"></td>
                <td style="padding: 4px 0; text-align: right;">La Personne Responsable des Marchés</td>
            </tr>
        </table>
        <div style="margin: 20px 0;"><strong>Objet : LETTRE DE MANIFESTATION D'INTERET</strong></div>
        <div style="margin: 20px 0; text-align: justify;">
        <p>Nous, L'ENTREPRISE <strong>{entreprise.raisonSociale}</strong>, montrons notre intérêt face à {description} dans les domaines suivants :</p>
        <p style="margin: 15px 0; font-weight: bold;">PRESTATIONS DE SERVICES</p>
        <p style="margin: 10px 0; padding-left: 20px;">{domaines}</p>
        <p style="margin-top: 20px;">Dans l'espoir de retenir votre attention lors de vos consultations, nous vous prions de croire, en l'assurance de nos sentiments les plus dévoués.</p>
        </div>
        <div style="margin-top: 50px;">
        <div style="float: right; text-align: right;">
        Ouagadougou, le {date_actuelle}<br><br>Le responsable<br><br><strong>{responsable}</strong>
        </div>
        <div style="clear: both;"></div>
        </div>
        </div>"""

def construire_html_presentation(request, entreprise, description):
    """Construit le HTML pour la présentation de l'entreprise"""
    
    entete_image_html = get_entete_image_html(request, entreprise)
   
    return f"""<div style="font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto;">
        {entete_image_html}
        <h2 style="text-align: center; margin-bottom: 20px;">PRÉSENTATION DE L'ENTREPRISE</h2>
        <div style="margin-bottom: 15px;">
        <strong>Raison sociale :</strong> {entreprise.raisonSociale}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Domaine d'activité :</strong> {entreprise.domaineActive or 'Non renseigné'}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Localisation :</strong> {entreprise.localisation or 'Non renseignée'}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Année de création :</strong> {entreprise.annee_creation or 'Non renseignée'}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Taille :</strong> {entreprise.taille or 'Non renseignée'} employés
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Chiffre d'affaires :</strong> {entreprise.chiffre_affaires or 'Non renseigné'} FCFA
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Capital social :</strong> {entreprise.capital_social or 'Non renseigné'} FCFA
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Contact :</strong><br>
        Tél : {entreprise.telephone or 'Non renseigné'}<br>
        Email : {entreprise.email or 'Non renseigné'}<br>
        Site web : {entreprise.site_web or 'Non renseigné'}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Compétences clés :</strong><br>
        {entreprise.competencesCles or 'Non renseignées'}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>Présentation :</strong><br>
        {entreprise.description or 'Aucune description fournie'}
        </div>
        </div>"""

def construire_html_fiche(request, entreprise):
    """Construit le HTML pour la fiche de renseignement"""
    
    entete_image_html = get_entete_image_html(request, entreprise)
    
    # Récupérer les informations
    raison_sociale = entreprise.raisonSociale
    adresse = entreprise.localisation or "Non renseignée"
    telephone = str(entreprise.telephone) if entreprise.telephone else "Non renseigné"
    email = entreprise.email or "Non renseigné"
    rccm = getattr(entreprise, 'rccm', 'Non renseigné')
    ifu = getattr(entreprise, 'ifu', 'Non renseigné')
    domaine_activite = entreprise.domaineActive or "Non renseigné"
    responsable = getattr(entreprise, 'responsable_nom', 'Non renseigné')
    
    return f"""<div style="font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto;">
        {entete_image_html}
        <h2 style="text-align: center; margin-bottom: 30px;">FICHE DE RENSEIGNEMENT</h2>
        <div style="margin-bottom: 15px;">
        <strong>1. Nom ou raison sociale :</strong> {raison_sociale}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>2. Adresse :</strong> {adresse}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>3. Téléphone :</strong> {telephone}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>4. E-mail :</strong> {email}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>5. RCCM :</strong> {rccm}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>6. IFU :</strong> {ifu}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>7. Principale activité :</strong> {domaine_activite}
        </div>
        <div style="margin-bottom: 15px;">
        <strong>8. Personne responsable :</strong> {responsable}
        </div>
        </div>"""

def construire_html_materiel(request, entreprise):
    """Construit le HTML pour la liste du matériel"""
    
    entete_image_html = get_entete_image_html(request, entreprise)
    
    # Récupérer la liste du matériel
    materiels = MaterielEntreprise.objects.filter(entreprise=entreprise)
    
    # Générer les lignes du tableau
    lignes = ""
    for i, m in enumerate(materiels, 1):
        lignes += f"""
        <tr style="border-bottom: 1px solid #ccc;">
            <td style="padding: 8px; text-align: center;">{i:02d}</td>
            <td style="padding: 8px;">{m.designation}</td>
            <td style="padding: 8px; text-align: center;">{m.quantite}</td>
        </tr>
        """
    
    if not lignes:
        lignes = '<tr><td colspan="3" style="padding: 20px; text-align: center;">Aucun matériel renseigné</td></tr>'
    
    date_actuelle = timezone.now().strftime('%d/%m/%Y')
    responsable = request.user.get_full_name() or request.user.username or entreprise.raisonSociale
    
    return f"""<div style="font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto;">
        {entete_image_html}
        <h2 style="text-align: center; margin-bottom: 30px;">LISTE DE MATÉRIELS</h2>
        <table style="width: 100%; border-collapse: collapse; border: 1px solid #000;">
            <thead>
                <tr style="background-color: #f0f0f0; border-bottom: 1px solid #000;">
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">N°</th>
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">Type, capacité et caractéristiques du matériel</th>
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">Nombre</th>
                </tr>
            </thead>
            <tbody>
                {lignes}
            </tbody>
        </table>
        <div style="margin-top: 40px;">
        <div style="float: right; text-align: right;">
        Ouagadougou, le {date_actuelle}<br><br>
        Le responsable<br><br>
        <strong>{responsable}</strong>
        </div>
        <div style="clear: both;"></div>
        </div>
        </div>"""

def construire_html_personnel(request, entreprise):
    """Construit le HTML pour la liste du personnel"""
    
    entete_image_html = get_entete_image_html(request, entreprise)
   
    personnels = PersonnelCle.objects.filter(entreprise=entreprise)
    
    lignes = ""
    for i, p in enumerate(personnels, 1):
        lignes += f"""
        <tr style="border-bottom: 1px solid #ccc;">
            <td style="padding: 8px; text-align: center;">{i:02d}</td>
            <td style="padding: 8px;">{p.nom_prenom}</td>
            <td style="padding: 8px;">{p.poste}</td>
            <td style="padding: 8px;">{p.qualification}</td>
            <td style="padding: 8px; text-align: center;">{p.annees_experience}</td>
        </tr>
        """
    
    if not lignes:
        lignes = '<tr><td colspan="5" style="padding: 20px; text-align: center;">Aucun personnel renseigné</td></tr>'
    
    date_actuelle = timezone.now().strftime('%d/%m/%Y')
    responsable = request.user.get_full_name() or request.user.username or entreprise.raisonSociale
    
    return f"""<div style="font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto;">
        {entete_image_html}
        <h2 style="text-align: center; margin-bottom: 30px;">LISTE DU PERSONNEL CADRE</h2>
        <table style="width: 100%; border-collapse: collapse; border: 1px solid #000;">
            <thead>
                <tr style="background-color: #f0f0f0; border-bottom: 1px solid #000;">
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">N°</th>
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">Nom et prénom</th>
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">Poste</th>
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">Qualification</th>
                    <th style="padding: 10px; text-align: center; border: 1px solid #000;">Années d'exp.</th>
                </tr>
            </thead>
            <tbody>
                {lignes}
            </tbody>
        </table>
        <div style="margin-top: 40px;">
        <div style="float: right; text-align: right;">
        Ouagadougou, le {date_actuelle}<br><br>
        Le responsable<br><br>
        <strong>{responsable}</strong>
        </div>
        <div style="clear: both;"></div>
        </div>
        </div>"""

def get_pdf_bytes(html_content):
    """Retourne les bytes d'un PDF (sans FileResponse)"""
    buffer = BytesIO()
    HTML(string=html_content).write_pdf(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def get_docx_bytes(html_content):
    """Retourne les bytes d'un DOCX à partir du HTML"""
    buffer = BytesIO()
    
    # Extraire le texte brut du HTML
    text = re.sub(r'<[^>]+>', ' ', html_content)
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Créer le document Word
    doc = DocxDocument()
    doc.add_paragraph(text)
    
    # Sauvegarder dans le buffer
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def generer_message_enveloppe(entreprise, opportunite):
    """
    Génère le message à coller sur l'enveloppe
    Contenu: RÉPONSE + description de l'AMI (en majuscules, en gras)
    """
    # Récupérer la description de l'AMI
    description = opportunite.description or ""
    
    # Nettoyer la description
    description = description.replace('\n', ' ').replace('\r', ' ')
    description = description.strip()
    
    # Supprimer les préfixes
    description = re.sub(r'Avis à manifestation d[’\']intérêt\s+', '', description, flags=re.IGNORECASE)
    description = re.sub(r'AVIS À MANIFESTATION D[’\']INTÉRÊT\s+', '', description, flags=re.IGNORECASE)
    
    # Mettre en majuscules
    description = description.upper()
    
    # Utiliser la fonction centralisée
    texte_brut = f"RÉPONSE {description}"
    return construire_html_enveloppe(texte_brut)

def generer_pdf(html_content, filename):
    return FileResponse(
        BytesIO(get_pdf_bytes(html_content)),  # ← réutilise get_pdf_bytes
        as_attachment=True,
        filename=filename
    )

def generer_docx(html_content, filename):
    """Génère un DOCX à partir du HTML"""
    return FileResponse(
        BytesIO(get_docx_bytes(html_content)),
        as_attachment=True,
        filename=filename,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

def generer_lettre_manifestation(request, entreprise, opportunite):
    """Génère la lettre de manifestation d'intérêt"""
    
    description = opportunite.description or ""
    description = description.strip()
    
    domaines = entreprise.domaineActive or "Fournitures et équipements ; Services courants ; Prestations intellectuelles"
    date_actuelle = timezone.now().strftime('%d/%m/%Y')
    responsable = request.user.get_full_name() or request.user.username or entreprise.raisonSociale
    
    return construire_html_lettre(request, entreprise, description, responsable, date_actuelle, domaines)

def generer_presentation_entreprise(request, entreprise):
    """Génère la présentation de l'entreprise"""
    return construire_html_presentation(request, entreprise, entreprise.description or "")

def generer_fiche_renseignement(request, entreprise):
    """Génère la fiche de renseignement de l'entreprise"""
    return construire_html_fiche(request, entreprise)

def generer_liste_materiel(request, entreprise):
    """Génère la liste du matériel"""
    return construire_html_materiel(request, entreprise)

def generer_liste_personnel(request, entreprise):
    """Génère la liste du matériel"""
    return construire_html_personnel(request, entreprise)

def generer_zip_documents(dossier, format_doc):
    """
    Génère un ZIP contenant tous les documents du dossier
    
    Args:
        dossier: DossierSoumission
        format_doc: 'pdf' ou 'docx'
    
    Returns:
        BytesIO: Buffer contenant le ZIP
    """
    noms_fichiers = {
        'ENVELOPPE': '01_message_enveloppe',
        'LETTRE': '02_lettre_manifestation',
        'PRESENTATION': '03_presentation_entreprise',
        'FICHE': '04_fiche_renseignement',
        'MATERIEL': '05_liste_materiel',
        'PERSONNEL': '06_liste_personnel',
    }
    
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for doc in dossier.documents_prepares.all():
            if doc.contenu_html and doc.statut in ['GENERATED', 'IMPORTED', 'MODIFIED']:
                nom = noms_fichiers.get(doc.type_document, doc.type_document.lower())
                filename = f"{nom}.{format_doc}"
                
                if format_doc == 'pdf':
                    file_content = get_pdf_bytes(doc.contenu_html)
                else:
                    file_content = get_docx_bytes(doc.contenu_html)
                
                zip_file.writestr(filename, file_content)
    
    zip_buffer.seek(0)
    return zip_buffer

@login_required
def api_document_generer(request):
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
        print(f"✅ JSON décodé: {data}")
    except json.JSONDecodeError as e:
        print(f"❌ Erreur JSON: {e}")
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    opportunite_type = data.get('opportunite_type')
    opportunite_id = data.get('opportunite_id')
    type_document = data.get('type_document')
    
    # Vérifier chaque paramètre individuellement
    if not opportunite_type:
        print("❌ opportunite_type manquant")
        return JsonResponse({'error': 'opportunite_type manquant'}, status=400)
    
    if not opportunite_id:
        print("❌ opportunite_id manquant")
        return JsonResponse({'error': 'opportunite_id manquant'}, status=400)
    
    if not type_document:
        print("❌ type_document manquant")
        return JsonResponse({'error': 'type_document manquant'}, status=400)
    
    print("✅ Tous les paramètres sont présents")
    
    # Convertir en majuscules pour la base de données
    type_document_upper = type_document.upper()
    print(f"🔄 type_document converti: {type_document} -> {type_document_upper}")
    
    # Vérifier que c'est un AMI
    if opportunite_type != 'Ami_uemoa':
        print(f"❌ Type non supporté: {opportunite_type}")
        return JsonResponse({'error': 'Seuls les AMI sont supportés'}, status=400)
    
    # Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
        print(f"✅ Entreprise trouvée: {entreprise.raisonSociale}")
    except Entreprise.DoesNotExist:
        print("❌ Entreprise non trouvée")
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Récupérer l'opportunité
    try:
        opportunite = Ami_uemoa.objects.get(id=opportunite_id)
        print(f"✅ Opportunité trouvée: AMI #{opportunite_id}")
    except Ami_uemoa.DoesNotExist:
        print(f"❌ AMI #{opportunite_id} non trouvé")
        return JsonResponse({'error': 'AMI non trouvé'}, status=404)
    
    # Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
        print(f"✅ Dossier trouvé: #{dossier.id}")
    except DossierSoumission.DoesNotExist:
        print("❌ Dossier non trouvé")
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # Vérifier que le dossier n'est pas déjà soumis
    if dossier.statut == 'SOUMIS':
        print("❌ Dossier déjà soumis")
        return JsonResponse({'error': 'Dossier déjà soumis, modification impossible'}, status=400)
    
    # Récupérer ou créer le document
    doc, created = DocumentGenere.objects.get_or_create(
        dossier=dossier,
        type_document=type_document_upper,
        defaults={'statut': 'MISSING', 'version': 1}
    )
    print(f"📄 Document: {'créé' if created else 'existait déjà'}")
    
    # Génération selon le type
    if type_document == 'enveloppe':
        print("🎨 Génération du message enveloppe...")
        contenu = generer_message_enveloppe(entreprise, opportunite)
        doc.contenu_html = contenu
        doc.statut = 'GENERATED'
        doc.save()
        print("✅ Message enveloppe généré avec succès")
        return JsonResponse({'success': True, 'message': 'Message enveloppe généré avec succès'})
    
    elif type_document == 'lettre':
        contenu = generer_lettre_manifestation(request, entreprise, opportunite)
        doc.contenu_html = contenu
        doc.statut = 'GENERATED'
        doc.save()
        return JsonResponse({'success': True, 'message': 'Lettre générée avec succès'})
    
    elif type_document == 'presentation':
        contenu = generer_presentation_entreprise(request, entreprise)
        doc.contenu_html = contenu
        doc.statut = 'GENERATED'
        doc.save()
        return JsonResponse({'success': True, 'message': 'Présentation générée avec succès'})
    
    elif type_document == 'fiche':
        contenu = generer_fiche_renseignement(request, entreprise)
        doc.contenu_html = contenu
        doc.statut = 'GENERATED'
        doc.save()
        return JsonResponse({'success': True, 'message': 'Fiche de renseignement générée avec succès'})
    
    elif type_document == 'materiel':
        contenu = generer_liste_materiel(request, entreprise)
        doc.contenu_html = contenu
        doc.statut = 'GENERATED'
        doc.save()
        return JsonResponse({'success': True, 'message': 'Liste du matériel générée avec succès'})
    
    elif type_document == 'personnel':
        contenu = generer_liste_personnel(request, entreprise)
        doc.contenu_html = contenu
        doc.statut = 'GENERATED'
        doc.save()
        return JsonResponse({'success': True, 'message': 'Liste du personnel générée avec succès'})
    
    else:
        print(f"❌ Type de document non supporté: {type_document}")
        return JsonResponse({'error': f'Type de document non supporté: {type_document}'}, status=400) 

@login_required
def api_document_apercu(request):
    """
    GET /api/document/apercu/?opportunite_type=Ami_uemoa&opportunite_id=66&type_document=enveloppe
    Retourne le HTML du document pour aperçu
    """
    if request.method != 'GET':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    opportunite_type = request.GET.get('opportunite_type')
    opportunite_id = request.GET.get('opportunite_id')
    type_document = request.GET.get('type_document')
    
    if not all([opportunite_type, opportunite_id, type_document]):
        return JsonResponse({'error': 'Paramètres manquants'}, status=400)
    
    # Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
    except DossierSoumission.DoesNotExist:
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # Récupérer le document
    type_document_upper = type_document.upper()
    try:
        doc = DocumentGenere.objects.get(
            dossier=dossier,
            type_document=type_document_upper
        )
    except DocumentGenere.DoesNotExist:
        return JsonResponse({'error': 'Document non trouvé'}, status=404)
    
    # Retourner le contenu HTML
    if doc.contenu_html:
        print("="*50)
        print(f"APERÇU LETTRE - Premier 500 caractères:")
        print(doc.contenu_html[:500])
        print("="*50)
        return JsonResponse({'success': True, 'html': doc.contenu_html})
    else:
        return JsonResponse({'error': 'Aucun contenu disponible pour ce document'}, status=404)

@login_required
def api_document_modifier(request):
    """
    POST /api/document/modifier/
    Modifie le contenu d'un document existant
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    opportunite_type = data.get('opportunite_type')
    opportunite_id = data.get('opportunite_id')
    type_document = data.get('type_document')
    nouveau_texte = data.get('contenu')
    
    if not all([opportunite_type, opportunite_id, type_document, nouveau_texte]):
        return JsonResponse({'error': 'Paramètres manquants'}, status=400)
    
    # Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
    except DossierSoumission.DoesNotExist:
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # Vérifier que le dossier n'est pas déjà soumis
    if dossier.statut == 'SOUMIS':
        return JsonResponse({'error': 'Dossier déjà soumis, modification impossible'}, status=400)
    
    # Récupérer le document
    type_document_upper = type_document.upper()
    try:
        doc = DocumentGenere.objects.get(
            dossier=dossier,
            type_document=type_document_upper
        )
    except DocumentGenere.DoesNotExist:
        return JsonResponse({'error': 'Document non trouvé'}, status=404)
    
    # ============================================
    # RECONSTRUIRE LE HTML SELON LE TYPE
    # ============================================
    
    if type_document == 'enveloppe':
        nouveau_html = construire_html_enveloppe(nouveau_texte)
    
    elif type_document == 'lettre':
        opportunite = get_object_or_404(Ami_uemoa, id=opportunite_id)
        date_actuelle = timezone.now().strftime('%d/%m/%Y')
        responsable = request.user.get_full_name() or request.user.username or entreprise.raisonSociale
        domaines = entreprise.domaineActive or "Fournitures et équipements ; Services courants ; Prestations intellectuelles"
        nouveau_html = construire_html_lettre(entreprise, nouveau_texte, responsable, date_actuelle, domaines)
    
    elif type_document == 'presentation':
        nouveau_html = construire_html_presentation(entreprise, nouveau_texte)
    
    elif type_document == 'fiche':
        nouveau_html = construire_html_fiche(entreprise)
    
    elif type_document == 'materiel':
        nouveau_html = construire_html_materiel(request, entreprise)
        
    else:
        nouveau_html = nouveau_texte
    
    # Mettre à jour le contenu (UNE SEULE FOIS)
    doc.contenu_html = nouveau_html
    doc.statut = 'MODIFIED'
    doc.version += 1
    doc.save()
    
    return JsonResponse({'success': True, 'message': 'Document modifié avec succès'})

@login_required
def api_document_telecharger(request):
    """
    GET /api/document/telecharger/?opportunite_type=Ami_uemoa&opportunite_id=66&type_document=enveloppe&format=pdf
    """
    if request.method != 'GET':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    opportunite_type = request.GET.get('opportunite_type')
    opportunite_id = request.GET.get('opportunite_id')
    type_document = request.GET.get('type_document')
    format_doc = request.GET.get('format', 'pdf')
    
    if not all([opportunite_type, opportunite_id, type_document]):
        return JsonResponse({'error': 'Paramètres manquants'}, status=400)
    
    # Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
    except DossierSoumission.DoesNotExist:
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # Récupérer le document
    type_document_upper = type_document.upper()
    try:
        doc = DocumentGenere.objects.get(
            dossier=dossier,
            type_document=type_document_upper
        )
    except DocumentGenere.DoesNotExist:
        return JsonResponse({'error': 'Document non trouvé'}, status=404)
    
    if not doc.contenu_html:
        return JsonResponse({'error': 'Aucun contenu généré'}, status=404)
    
    # Générer le fichier selon le format
    filename = f"{type_document}_{dossier.id}_{timezone.now().strftime('%Y%m%d')}"
    
    if format_doc == 'pdf':
        return generer_pdf(doc.contenu_html, f"{filename}.pdf")
    elif format_doc == 'docx':
        return generer_docx(doc.contenu_html, f"{filename}.docx")
    else:
        return JsonResponse({'error': 'Format non supporté'}, status=400)

def convertir_pdf_en_html(fichier):
    """Convertit un PDF en HTML (version simplifiée)"""
    # Pour l'instant, extraire le texte du PDF
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(fichier)
        texte = ""
        for page in pdf_reader.pages:
            texte += page.extract_text()
        
        # Encapsuler dans du HTML
        html = f"""
        <div style="font-family: Arial, sans-serif; padding: 20px;">
            <div style="white-space: pre-wrap;">
                {texte}
            </div>
        </div>
        """
        return html
    except ImportError:
        return "<div>Import PDF - Installation de PyPDF2 requise</div>"

def convertir_docx_en_html(fichier):
    """Convertit un DOCX en HTML"""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(fichier)
        texte = ""
        for paragraph in doc.paragraphs:
            texte += paragraph.text + "\n"
        
        html = f"""
        <div style="font-family: Arial, sans-serif; padding: 20px;">
            <div style="white-space: pre-wrap;">
                {texte}
            </div>
        </div>
        """
        return html
    except ImportError:
        return "<div>Import DOCX - Installation de python-docx requise</div>"

@login_required
def api_document_importer(request):
    """
    POST /api/document/importer/
    Importe un fichier (PDF ou DOCX) et le convertit en HTML
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    opportunite_type = request.POST.get('opportunite_type')
    opportunite_id = request.POST.get('opportunite_id')
    type_document = request.POST.get('type_document')
    fichier = request.FILES.get('fichier')
    
    if not all([opportunite_type, opportunite_id, type_document, fichier]):
        return JsonResponse({'error': 'Paramètres manquants'}, status=400)
    
    # Vérifier l'extension du fichier
    nom_fichier = fichier.name
    extension = nom_fichier.split('.')[-1].lower()
    
    if extension not in ['pdf', 'docx']:
        return JsonResponse({'error': 'Format non supporté. Utilisez PDF ou DOCX'}, status=400)
    
    # Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
    except DossierSoumission.DoesNotExist:
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    if dossier.statut == 'SOUMIS':
        return JsonResponse({'error': 'Dossier déjà soumis'}, status=400)
    
    # Récupérer le document
    type_document_upper = type_document.upper()
    try:
        doc = DocumentGenere.objects.get(
            dossier=dossier,
            type_document=type_document_upper
        )
    except DocumentGenere.DoesNotExist:
        return JsonResponse({'error': 'Document non trouvé'}, status=404)
    
    # Convertir le fichier en HTML
    try:
        if extension == 'pdf':
            contenu_html = convertir_pdf_en_html(fichier)
        else:  # docx
            contenu_html = convertir_docx_en_html(fichier)
    except Exception as e:
        return JsonResponse({'error': f'Erreur lors de la conversion: {str(e)}'}, status=500)
    
    # Sauvegarder
    doc.contenu_html = contenu_html
    doc.statut = 'IMPORTED'
    doc.version += 1
    doc.save()
    
    return JsonResponse({'success': True, 'message': 'Document importé avec succès'})

@login_required
def api_document_supprimer(request):
    """API : supprimer (archiver) un document"""
    pass

@login_required
def api_dossier_soumettre(request):
    """
    POST /api/dossier/soumettre/
    Soumet le dossier définitivement et envoie par email
    """

    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
        print(f"✅ JSON décodé: {data}")
    except json.JSONDecodeError:
        print(f"❌ Erreur JSON: {e}")
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    opportunite_type = data.get('opportunite_type')
    opportunite_id = data.get('opportunite_id')
    email_destinataire = data.get('email_destinataire', '')
    format_doc = data.get('format', 'pdf')

    if not opportunite_type or not opportunite_id:
        print("❌ opportunite_type ou opportunite_id manquant")
        return JsonResponse({'error': 'Paramètres manquants'}, status=400)
    
    # Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Récupérer l'opportunité
    if opportunite_type == 'Ami_uemoa':
        opportunite = get_object_or_404(Ami_uemoa, id=opportunite_id)
    else:
        opportunite = None
    
    # Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
    except DossierSoumission.DoesNotExist:
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # Vérifier que tous les documents sont prêts
    documents = dossier.documents_prepares.all()
    tous_ok = all(doc.statut in ['GENERATED', 'IMPORTED', 'MODIFIED'] for doc in documents)
    
    if not tous_ok:
        return JsonResponse({'error': 'Tous les documents doivent être générés avant soumission'}, status=400)
    
    # Vérifier que le dossier n'est pas déjà soumis
    if dossier.statut == 'SOUMIS':
        return JsonResponse({'error': 'Ce dossier a déjà été soumis'}, status=400)
    
    # ============================================
    # DÉTERMINER L'EMAIL DU DESTINATAIRE
    # ============================================
    
    if not email_destinataire:
        if opportunite:
            if hasattr(opportunite, 'email_contact') and opportunite.email_contact:
                email_destinataire = opportunite.email_contact
            elif hasattr(opportunite, 'email') and opportunite.email:
                email_destinataire = opportunite.email
    
    if not email_destinataire:
        return JsonResponse({'error': 'Email destinataire requis. Veuillez le saisir.'}, status=400)
    
    # ============================================
    # GÉNÉRER LE ZIP
    # ============================================
    
    try:
        zip_buffer = generer_zip_documents(dossier, format_doc)
    except Exception as e:
        return JsonResponse({'error': f'Erreur lors de la génération du ZIP: {str(e)}'}, status=500)
    
    # ============================================
    # EMAIL AU DESTINATAIRE
    # ============================================
    
    sujet = f"Dossier de candidature - {entreprise.raisonSociale} - {dossier.reference}"
    message = f"""
Bonjour,

L'entreprise {entreprise.raisonSociale} a soumis son dossier de candidature.

Référence du dossier : {dossier.reference}
Date de soumission : {timezone.now().strftime('%d/%m/%Y à %H:%M')}
Format du dossier : {format_doc.upper()}

Vous trouverez le dossier complet en pièce jointe.

Cordialement,
L'équipe FASOIA
"""
    
    email = EmailMessage(
        subject=sujet,
        body=message,
        to=[email_destinataire],
        reply_to=[entreprise.email],
    )
    
    zip_buffer.seek(0)
    email.attach(f"dossier_complet_{dossier.reference}.zip", zip_buffer.read(), 'application/zip')
    
    # ============================================
    # EMAIL DE COPIE À L'ENTREPRISE
    # ============================================
    
    message_copy = f"""
Bonjour {entreprise.raisonSociale},

Vous trouverez ci-joint une copie du dossier que vous avez soumis.

Référence : {dossier.reference}
Date de soumission : {timezone.now().strftime('%d/%m/%Y à %H:%M')}
Destinataire : {email_destinataire}
Format : {format_doc.upper()}

Cordialement,
L'équipe FASOIA
"""
    
    email_copy = EmailMessage(
        subject=f"Copie de votre soumission - {dossier.reference}",
        body=message_copy,
        to=[entreprise.email],
    )
    
    # Re-générer le ZIP pour la copie
    zip_buffer2 = generer_zip_documents(dossier, format_doc)
    email_copy.attach(f"dossier_complet_{dossier.reference}.zip", zip_buffer2.read(), 'application/zip')
    
    # ============================================
    # ENVOYER LES EMAILS
    # ============================================
    
    try:
        email.send()
        email_copy.send()
    except Exception as e:
        return JsonResponse({'error': f'Erreur lors de l\'envoi des emails: {str(e)}'}, status=500)
    
    # ============================================
    # MARQUER COMME SOUMIS
    # ============================================
    
    dossier.statut = 'SOUMIS'
    dossier.date_soumission_effective = timezone.now()
    dossier.save()
    
    # Mettre à jour les statistiques
    entreprise.nb_candidatures_emises += 1
    entreprise.save()
    
    return JsonResponse({
        'success': True,
        'message': f'Dossier soumis avec succès et envoyé à {email_destinataire}',
        'date_soumission': dossier.date_soumission_effective.strftime('%d/%m/%Y à %H:%M')
    })

@login_required
def api_dossier_etat(request):
    """
    GET /api/dossier/etat/?opportunite_type=Ami_uemoa&opportunite_id=66
    Retourne l'état complet d'un dossier de soumission
    """
    if request.method != 'GET':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    # 1. Récupérer les paramètres
    opportunite_type = request.GET.get('opportunite_type')
    opportunite_id = request.GET.get('opportunite_id')
    
    if not opportunite_type or not opportunite_id:
        return JsonResponse({'error': 'Paramètres manquants'}, status=400)
    
    # 2. Vérifier que c'est un AMI (pour l'instant)
    if opportunite_type != 'Ami_uemoa':
        return JsonResponse({'error': 'Seuls les AMI sont supportés pour le moment'}, status=400)
    
    # 3. Récupérer l'entreprise
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # 4. Récupérer le dossier
    try:
        dossier = DossierSoumission.objects.get(
            entreprise=entreprise,
            opportunite_type=opportunite_type,
            opportunite_id=opportunite_id
        )
    except DossierSoumission.DoesNotExist:
        return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # 5. Récupérer tous les documents du dossier
    documents = dossier.documents_prepares.all()
    
    # 6. Construire la réponse
    etat = {
        'dossier_id': dossier.id,
        'statut': dossier.statut,
        'est_soumis': dossier.statut == 'SOUMIS',
        'documents': {}
    }
    
    # Mapping des types de document (DB -> frontend)
    mapping = {
        'ENVELOPPE': 'enveloppe',
        'LETTRE': 'lettre',
        'PRESENTATION': 'presentation',
        'FICHE': 'fiche',
        'MATERIEL': 'materiel',
        'PERSONNEL': 'personnel',
    }
    
    for doc in documents:
        type_key = mapping.get(doc.type_document, doc.type_document.lower())
        etat['documents'][type_key] = {
            'statut': doc.statut,
            'version': doc.version,
            'date_generation': doc.date_generation.isoformat() if doc.date_generation else None,
            'date_modification': doc.date_modification.isoformat() if doc.date_modification else None,
        }
    
    return JsonResponse({'success': True, 'data': etat})

@login_required
def api_dossier_telecharger_complet(request, dossier_id=None):
    """
    GET /api/dossier/telecharger-complet/?opportunite_type=...&opportunite_id=...
    GET /api/dossier/telecharger-complet/<dossier_id>/
    """
    if request.method != 'GET':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    format_doc = request.GET.get('format', 'pdf')
    
    if format_doc not in ['pdf', 'docx']:
        return JsonResponse({'error': 'Format non supporté'}, status=400)
    
    try:
        entreprise = Entreprise.objects.get(user=request.user)
    except Entreprise.DoesNotExist:
        return JsonResponse({'error': 'Entreprise non trouvée'}, status=404)
    
    # Cas 1: Téléchargement par ID
    if dossier_id:
        try:
            dossier = DossierSoumission.objects.get(id=dossier_id, entreprise=entreprise)
        except DossierSoumission.DoesNotExist:
            return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    else:
        # Cas 2: Téléchargement par opportunite_type + opportunite_id
        opportunite_type = request.GET.get('opportunite_type')
        opportunite_id = request.GET.get('opportunite_id')
        
        if not opportunite_type or not opportunite_id:
            return JsonResponse({'error': 'Paramètres manquants'}, status=400)
        
        try:
            dossier = DossierSoumission.objects.get(
                entreprise=entreprise,
                opportunite_type=opportunite_type,
                opportunite_id=opportunite_id
            )
        except DossierSoumission.DoesNotExist:
            return JsonResponse({'error': 'Dossier non trouvé'}, status=404)
    
    # Générer le ZIP
    try:
        zip_buffer = generer_zip_documents(dossier, format_doc)
    except Exception as e:
        return JsonResponse({'error': f'Erreur lors de la génération: {str(e)}'}, status=500)
    
    return FileResponse(
        zip_buffer,
        as_attachment=True,
        filename=f"dossier_complet_{dossier.reference}_{timezone.now().strftime('%Y%m%d')}.zip",
        content_type='application/zip'
    )

@login_required
def api_materiel_liste(request):
    """Retourne la liste du matériel de l'entreprise"""
    entreprise = Entreprise.objects.get(user=request.user)
    materiels = MaterielEntreprise.objects.filter(entreprise=entreprise)
    
    data = [
        {
            'id': m.id,
            'designation': m.designation,
            'quantite': m.quantite,
        }
        for m in materiels
    ]
    return JsonResponse({'success': True, 'materiels': data})

@login_required
def api_materiel_ajouter(request):
    """Ajoute un matériel"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    entreprise = Entreprise.objects.get(user=request.user)
    
    materiel = MaterielEntreprise.objects.create(
        entreprise=entreprise,
        designation=data.get('designation'),
        quantite=data.get('quantite', 1)
    )
    
    return JsonResponse({
        'success': True, 
        'materiel': {'id': materiel.id, 'designation': materiel.designation, 'quantite': materiel.quantite}
    })

@login_required
def api_materiel_modifier(request):
    """Modifie un matériel existant"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    entreprise = Entreprise.objects.get(user=request.user)
    
    try:
        materiel = MaterielEntreprise.objects.get(id=data.get('id'), entreprise=entreprise)
        materiel.designation = data.get('designation', materiel.designation)
        materiel.quantite = data.get('quantite', materiel.quantite)
        materiel.save()
        
        return JsonResponse({'success': True, 'message': 'Matériel modifié avec succès'})
    except MaterielEntreprise.DoesNotExist:
        return JsonResponse({'error': 'Matériel non trouvé'}, status=404)

@login_required
def api_materiel_supprimer(request):
    """Supprime un matériel"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    entreprise = Entreprise.objects.get(user=request.user)
    
    try:
        materiel = MaterielEntreprise.objects.get(id=data.get('id'), entreprise=entreprise)
        materiel.delete()
        return JsonResponse({'success': True, 'message': 'Matériel supprimé avec succès'})
    except MaterielEntreprise.DoesNotExist:
        return JsonResponse({'error': 'Matériel non trouvé'}, status=404)

@login_required
def api_personnel_liste(request):
    """Retourne la liste du personnel de l'entreprise"""
    entreprise = Entreprise.objects.get(user=request.user)
    personnels = PersonnelCle.objects.filter(entreprise=entreprise)
    
    data = [
        {
            'id': p.id,
            'nom_prenom': p.nom_prenom,
            'poste': p.poste,
            'qualification': p.qualification,
            'annees_experience': p.annees_experience,
        }
        for p in personnels
    ]
    return JsonResponse({'success': True, 'personnels': data})

@login_required
def api_personnel_ajouter(request):
    """Ajoute un personnel"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    entreprise = Entreprise.objects.get(user=request.user)
    
    personnel = PersonnelCle.objects.create(
        entreprise=entreprise,
        nom_prenom=data.get('nom_prenom'),
        poste=data.get('poste'),
        qualification=data.get('qualification'),
        annees_experience=data.get('annees_experience', 0)
    )
    
    return JsonResponse({
        'success': True,
        'personnel': {
            'id': personnel.id,
            'nom_prenom': personnel.nom_prenom,
            'poste': personnel.poste,
            'qualification': personnel.qualification,
            'annees_experience': personnel.annees_experience,
        }
    })

@login_required
def api_personnel_modifier(request):
    """Modifie un personnel existant"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    entreprise = Entreprise.objects.get(user=request.user)
    
    try:
        personnel = PersonnelCle.objects.get(id=data.get('id'), entreprise=entreprise)
        personnel.nom_prenom = data.get('nom_prenom', personnel.nom_prenom)
        personnel.poste = data.get('poste', personnel.poste)
        personnel.qualification = data.get('qualification', personnel.qualification)
        personnel.annees_experience = data.get('annees_experience', personnel.annees_experience)
        personnel.save()
        
        return JsonResponse({'success': True, 'message': 'Personnel modifié avec succès'})
    except PersonnelCle.DoesNotExist:
        return JsonResponse({'error': 'Personnel non trouvé'}, status=404)

@login_required
def api_personnel_supprimer(request):
    """Supprime un personnel"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'JSON invalide'}, status=400)
    
    entreprise = Entreprise.objects.get(user=request.user)
    
    try:
        personnel = PersonnelCle.objects.get(id=data.get('id'), entreprise=entreprise)
        personnel.delete()
        return JsonResponse({'success': True, 'message': 'Personnel supprimé avec succès'})
    except PersonnelCle.DoesNotExist:
        return JsonResponse({'error': 'Personnel non trouvé'}, status=404)

def trouver_emploi(request):
    """
    Page d'outils et de conseils pour la recherche d'emploi
    """
    # Récupérer les dernières offres d'emploi
    dernieres_offres = OffreEmploi.objects.filter(
        statut='PUBLIEE',
        est_active=True
    ).order_by('-date_publication')[:5]
    
    # Statistiques
    stats = {
        'offres_disponibles': OffreEmploi.objects.filter(
            statut='PUBLIEE', 
            est_active=True
        ).count(),
        'outils': 6,  # Nombre d'outils affichés dans la grille
        'conseils': 6,  # Nombre de conseils affichés
    }
    
    # Si l'utilisateur est un candidat, on peut personnaliser
    if hasattr(request.user, 'particulier'):
        particulier = request.user.particulier
        if hasattr(particulier, 'candidat'):
            candidat = particulier.candidat
            # On pourrait ajouter des recommandations personnalisées ici
            pass
    
    context = {
        'dernieres_offres': dernieres_offres,
        'stats': stats,
    }
    
    return render(request, 'myAppli/outils_emploi/trouver_emploi.html', context)

def generer_cv(request):
    """
    Générateur de CV accessible à tous.
    Gère la création, l'édition et la photo de profil.
    """

    modeles_cv  = ModeleCV.objects.filter(est_actif=True).order_by('ordre_affichage')
    modeles_json = json.dumps([
        {
            'nom':          m.nom,
            'categorie':    m.categorie,
            'est_populaire': m.est_populaire,
            'est_premium':  m.est_premium,
        }
        for m in modeles_cv
    ])

    # ── Mode édition ──────────────────────────────────────────────────────────
    cv_a_modifier = None
    donnees_init  = {}
    is_edit_mode  = False

    if request.GET.get('edit'):
        try:
            cv_id         = int(request.GET.get('edit'))
            cv_a_modifier = CVGenere.objects.get(id=cv_id, utilisateur=request.user)
            donnees_init  = cv_a_modifier.donnees_cv.copy()
            is_edit_mode  = True

            # Reconvertir les listes en chaînes pour le formulaire
            if isinstance(donnees_init.get('competences'), list):
                donnees_init['competences'] = ', '.join(donnees_init['competences'])
            if isinstance(donnees_init.get('langues'), list):
                donnees_init['langues'] = ', '.join(donnees_init['langues'])
            if isinstance(donnees_init.get('centres_interet'), list):
                donnees_init['centres_interet'] = ', '.join(donnees_init['centres_interet'])

            if isinstance(donnees_init.get('experiences'), list):
                exp_text = ''
                for exp in donnees_init['experiences']:
                    exp_text += (
                        f"{exp.get('titre', '')} | {exp.get('entreprise', '')} | "
                        f"{exp.get('date_debut', '')} | {exp.get('date_fin', 'Présent')} | "
                        f"{exp.get('description', '')}\n"
                    )
                donnees_init['experiences'] = exp_text.strip()

            if isinstance(donnees_init.get('formations'), list):
                form_text = ''
                for f in donnees_init['formations']:
                    form_text += f"{f.get('diplome', '')} | {f.get('etablissement', '')} | {f.get('annee', '')}\n"
                donnees_init['formations'] = form_text.strip()

            print(f"📝 Mode édition CV {cv_id} : {cv_a_modifier.titre}")

        except (CVGenere.DoesNotExist, ValueError) as e:
            print(f"Erreur mode édition : {e}")

    # ── Traitement POST ───────────────────────────────────────────────────────
    if request.method == 'POST':
        try:
            format_export = request.POST.get('format', 'pdf')
            style         = request.POST.get('style', 'executive')

            from .services.generateur_cv_public import GenerateurCVPublic
            generateur = GenerateurCVPublic()

            # ✅ On passe request.FILES pour récupérer la photo
            donnees = generateur.preparer_donnees(request.POST, request.FILES)

            # DEBUG
            print("=== VÉRIFICATION FINALE ===")
            print(f"Clés dans donnees: {donnees.keys()}")
            print(f"Formations: {donnees.get('formations')}")
            print(f"Type de formations: {type(donnees.get('formations'))}")
            print(f"Nombre de formations: {len(donnees.get('formations', []))}")
            
            if not donnees['prenom'] or not donnees['nom']:
                messages.error(request, "Le prénom et le nom sont obligatoires.")
                return redirect('myAppli:generer_cv')

            buffer, nom_fichier = generateur.generer_cv_buffer(donnees, format_export, style)

            # ── Mise à jour (mode édition) ────────────────────────────────────
            cv_id = request.POST.get('cv_id')
            if cv_id:
                try:
                    cv_existant             = CVGenere.objects.get(id=cv_id, utilisateur=request.user)
                    cv_existant.titre       = f"CV de {donnees['prenom']} {donnees['nom']}"
                    cv_existant.donnees_cv  = donnees
                    cv_existant.modele      = ModeleCV.objects.filter(categorie=style, est_actif=True).first()

                    buffer.seek(0)
                    if format_export == 'pdf':
                        if cv_existant.fichier_pdf:
                            cv_existant.fichier_pdf.delete()
                        cv_existant.fichier_pdf.save(nom_fichier, ContentFile(buffer.getvalue()))
                    elif format_export == 'docx':
                        if cv_existant.fichier_docx:
                            cv_existant.fichier_docx.delete()
                        cv_existant.fichier_docx.save(nom_fichier, ContentFile(buffer.getvalue()))

                    cv_existant.save()
                    messages.success(request, "✅ CV modifié avec succès !")
                    return redirect('myAppli:mes_cvs')

                except CVGenere.DoesNotExist:
                    messages.error(request, "CV non trouvé.")
                    return redirect('myAppli:generer_cv')

            # ── Création ──────────────────────────────────────────────────────
            else:
                if request.user.is_authenticated:
                    modele_cv  = ModeleCV.objects.filter(categorie=style, est_actif=True).first()
                    cv_nouveau = CVGenere.objects.create(
                        utilisateur=request.user,
                        modele=modele_cv,
                        titre=f"CV de {donnees['prenom']} {donnees['nom']}",
                        donnees_cv=donnees,
                        est_public=False,
                    )
                    buffer.seek(0)
                    if format_export == 'pdf':
                        cv_nouveau.fichier_pdf.save(nom_fichier, ContentFile(buffer.getvalue()))
                    elif format_export == 'docx':
                        cv_nouveau.fichier_docx.save(nom_fichier, ContentFile(buffer.getvalue()))

                    messages.success(request, "✅ CV créé avec succès !")
                    return redirect('myAppli:mes_cvs')
                else:
                    # Utilisateur non connecté → téléchargement direct
                    buffer.seek(0)
                    return FileResponse(buffer, as_attachment=True, filename=nom_fichier)

        except Exception as e:
            print(f"❌ Erreur génération CV : {e}")
            messages.error(request, f"Erreur lors de la génération : {str(e)}")
            return redirect('myAppli:generer_cv')

    # ── GET ───────────────────────────────────────────────────────────────────
    return render(request, 'myAppli/outils_emploi/generer_cv.html', {
        'modeles_cv':   modeles_cv,
        'modeles_json': modeles_json,
        'cv_a_modifier': cv_a_modifier,
        'donnees_init': donnees_init,
        'is_edit_mode': is_edit_mode,
    })

@login_required
def preparer_entretien(request):
    """
    Page principale — choisir offre ou cas général
    """
    if not hasattr(request.user, 'particulier'):
        return redirect('myAppli:home')

    candidat = request.user.particulier.candidat

    # Offres recommandées du candidat
    from analyse_ia.models import RecommandationEmploi
    offres_recommandees = OffreEmploi.objects.filter(
        recommandations__candidat=candidat
    ).order_by('-recommandations__score_global')[:10]

    # Sessions passées
    sessions = SessionEntretien.objects.filter(
        candidat=candidat
    ).order_by('-date_creation')[:5]

    return render(request, 'myAppli/entretien/preparer_entretien.html', {
        'candidat': candidat,
        'offres_recommandees': offres_recommandees,
        'sessions': sessions,
    })

@login_required
def demarrer_session(request):
    """
    Démarre une nouvelle session d'entretien
    et génère les 10 questions via Ollama
    """
    if request.method != 'POST':
        return redirect('myAppli:preparer_entretien')

    candidat = request.user.particulier.candidat
    offre_id = request.POST.get('offre_id')
    poste_vise = request.POST.get('poste_vise', '')

    # Récupérer l'offre si spécifique
    offre = None
    if offre_id:
        try:
            offre = OffreEmploi.objects.get(id=offre_id)
            poste_vise = offre.titre
            secteur = offre.get_secteur_display()
        except OffreEmploi.DoesNotExist:
            pass
    else:
        secteur = candidat.secteur_recherche

    # Construire le contexte pour Ollama
    if offre:
        contexte = f"""
        Poste: {offre.titre}
        Entreprise: {offre.entreprise_nom}
        Secteur: {offre.get_secteur_display()}
        Description: {offre.description[:500]}
        Compétences requises: {', '.join(offre.competences_requises)}
        Profil recherché: {offre.profil_recherche[:300]}
        """
    else:
        contexte = f"""
        Poste visé: {poste_vise}
        Secteur: {secteur}
        Compétences du candidat: {candidat.competences}
        Niveau d'étude: {candidat.niveauEtude}
        Années d'expérience: {candidat.anneesExperiences}
        """

    # Prompt pour générer 10 questions
    prompt = f"""
    Tu es un recruteur expert. Génère exactement 10 questions d'entretien
    pour le profil suivant :

    {contexte}

    Génère un mélange de questions :
    - 3 questions techniques (compétences spécifiques)
    - 3 questions comportementales (expériences passées)
    - 2 questions situationnelles (mises en situation)
    - 2 questions de motivation (pourquoi ce poste)

    Réponds UNIQUEMENT avec un JSON valide, sans texte avant ou après :
    {{
        "questions": [
            {{"ordre": 1, "type": "TECHNIQUE", "question": "..."}},
            {{"ordre": 2, "type": "COMPORTEMENTALE", "question": "..."}},
            ...
        ]
    }}
    """

    try:
        from analyse_ia.ia_client import IAClient
        ia = IAClient()
        questions_data = ia.generer_questions(contexte, poste_vise)

    except Exception as e:
        print(f"Erreur IA: {e}")
        messages.error(request, "Erreur lors de la génération des questions. Réessayez.")
        return redirect('myAppli:preparer_entretien')
    
    # Créer la session
    session = SessionEntretien.objects.create(
        candidat=candidat,
        offre=offre,
        poste_vise=poste_vise,
        secteur=secteur,
        statut=SessionEntretien.EN_COURS
    )

    # Créer les questions
    for q in questions_data[:10]:
        QuestionEntretien.objects.create(
            session=session,
            ordre=q.get('ordre', 1),
            type_question=q.get('type', 'COMPORTEMENTALE'),
            question=q.get('question', '')
        )

    messages.success(request, "Session démarrée ! Bonne chance 🎯")
    return redirect('myAppli:session_entretien', session_id=session.id)

@login_required
def session_entretien(request, session_id):
    """
    Page de la session — affiche question courante
    """
    candidat = request.user.particulier.candidat

    try:
        session = SessionEntretien.objects.get(
            id=session_id,
            candidat=candidat
        )
    except SessionEntretien.DoesNotExist:
        messages.error(request, "Session introuvable")
        return redirect('myAppli:preparer_entretien')

    # Question courante — première sans réponse
    question_courante = session.questions.filter(
        reponse_candidat=''
    ).first()

    derniere_question = session.questions.last()

    # Si toutes répondues → terminer
    if not question_courante and session.statut == SessionEntretien.EN_COURS:
        return redirect('myAppli:terminer_session', session_id=session.id)

    return render(request, 'myAppli/entretien/session_entretien.html', {
        'session': session,
        'question_courante': question_courante,
        'derniere_question' : derniere_question,
        'questions': session.questions.all(),
        'nb_repondues': session.nb_questions_repondues,
        'progression': int((session.nb_questions_repondues / 10) * 100),
    })

@login_required
def repondre_question(request, session_id, question_id):
    """
    Enregistre la réponse et demande feedback à Ollama
    """
    if request.method != 'POST':
        return redirect('myAppli:session_entretien', session_id=session_id)

    candidat = request.user.particulier.candidat
    reponse = request.POST.get('reponse', '').strip()

    if not reponse:
        messages.warning(request, "Veuillez saisir une réponse.")
        return redirect('myAppli:session_entretien', session_id=session_id)

    try:
        session = SessionEntretien.objects.get(id=session_id, candidat=candidat)
        question = QuestionEntretien.objects.get(id=question_id, session=session)
    except (SessionEntretien.DoesNotExist, QuestionEntretien.DoesNotExist):
        return redirect('myAppli:preparer_entretien')

    try:
        from analyse_ia.ia_client import IAClient
        ia = IAClient()
        evaluation = ia.evaluer_reponse(question, reponse, session.poste_vise)
        score = evaluation['score']
        feedback = evaluation['feedback']
        points_forts = evaluation['points_forts']
        points_amelioration = evaluation['points_amelioration']

    except Exception as e:
        print(f"Erreur évaluation: {e}")
        score = 5.0
        feedback = "Évaluation non disponible"
        points_forts = []
        points_amelioration = []
    
    # Sauvegarder la réponse
    question.reponse_candidat = reponse
    question.date_reponse = timezone.now()
    question.score = score
    question.feedback_ia = feedback
    question.points_forts = points_forts
    question.points_amelioration = points_amelioration
    question.save()

    return redirect('myAppli:session_entretien', session_id=session_id)

@login_required
def terminer_session(request, session_id):
    """
    Termine la session et génère le feedback global
    """

    candidat = request.user.particulier.candidat

    try:
        session = SessionEntretien.objects.get(id=session_id, candidat=candidat)
    except SessionEntretien.DoesNotExist:
        return redirect('myAppli:preparer_entretien')

    questions = session.questions.all()

    # Calculer score global
    scores = [q.score for q in questions if q.score > 0]
    score_global = round(sum(scores) / len(scores), 1) if scores else 0

    try:
        from analyse_ia.ia_client import IAClient
        ia = IAClient()
        feedback_global = ia.generer_feedback_global(session)
    except Exception as e:
        print(f"Erreur feedback global: {e}")
        feedback_global = "Session terminée avec succès !"

    # Mettre à jour la session
    session.statut = SessionEntretien.TERMINEE
    session.score_global = score_global
    session.feedback_global = feedback_global
    session.date_fin = timezone.now()
    session.save()

    return render(request, 'myAppli/entretien/resultat_entretien.html', {
        'session': session,
        'questions': questions,
        'score_global': score_global,
        'feedback_global': feedback_global,
    })

@login_required
def alertes_emploi(request):
    """
    Gestion des alertes emploi
    """
    return render(request, 'myAppli/outils_emploi /alertes_emploi.html')

def apercu_style_cv(request, style):
    """
    Génère un aperçu HTML d'un style de CV
    """
    generateur = GenerateurCVPublic()
    html_apercu = generateur.generer_apercu(style)
    return HttpResponse(html_apercu)

@login_required
def apercu_cv(request, cv_id):
    """Affiche un aperçu du CV en utilisant le template du style et les données sauvegardées"""
    cv = get_object_or_404(CVGenere, id=cv_id, utilisateur=request.user)
    
    # Récupérer les données du CV
    donnees = cv.donnees_cv
    
    # Récupérer le style
    style = donnees.get('style', 'moderne')
    if cv.modele:
        style = cv.modele.categorie
    
    print(f"🎨 Génération aperçu pour style: {style}")
    
    # Utiliser le générateur
    from .services.generateur_cv_public import GenerateurCVPublic
    generateur = GenerateurCVPublic()
    
    # Générer le HTML
    html = generateur.generer_html_avec_donnees(style, donnees)
    
    print(f"✅ HTML généré, longueur: {len(html)} caractères")
    
    return HttpResponse(html)

@login_required
def mes_cvs(request):
    """Affiche tous les CV de l'utilisateur"""
    cvs = CVGenere.objects.filter(utilisateur=request.user).order_by('-date_generation')
    
    # Statistiques avec gestion des champs optionnels
    stats = {
        'total': cvs.count(),
        'telechargements': sum(cv.nb_telechargements for cv in cvs if hasattr(cv, 'nb_telechargements')),
        'favoris': cvs.filter(est_favori=True).count() if hasattr(CVGenere, 'est_favori') else 0,
        'utilises': cvs.filter(est_utilise=True).count() if hasattr(CVGenere, 'est_utilise') else 0,
    }
    
    context = {
        'cvs': cvs,
        'stats': stats,
    }
    return render(request, 'myAppli/outils_emploi/mes_cvs.html', context)

@login_required
def telecharger_cv(request, cv_id, format):
    cv = get_object_or_404(CVGenere, id=cv_id, utilisateur=request.user)
    
    cv.nb_telechargements += 1
    cv.save()
    
    if format == 'pdf' and cv.fichier_pdf:
        return FileResponse(cv.fichier_pdf, as_attachment=True, filename=cv.fichier_pdf.name)
    elif format == 'docx' and cv.fichier_docx:
        return FileResponse(cv.fichier_docx, as_attachment=True, filename=cv.fichier_docx.name)
    
    return JsonResponse({'error': 'Fichier non disponible'}, status=404)

@login_required
def supprimer_cv(request, cv_id):
    cv = get_object_or_404(CVGenere, id=cv_id, utilisateur=request.user)
    cv.delete()
    return JsonResponse({'success': True})

@login_required
def dupliquer_cv(request, cv_id):
    original = get_object_or_404(CVGenere, id=cv_id, utilisateur=request.user)
    
    copie = CVGenere.objects.create(
        utilisateur=request.user,
        modele=original.modele,
        titre=f"{original.titre} (copie)",
        donnees_cv=original.donnees_cv,
        est_public=False
    )
    
    return JsonResponse({'success': True, 'id': copie.id})

@login_required
def importer_cv(request):
    """Importe un CV depuis un fichier"""
    print("="*50)
    print("🚀 FONCTION IMPORTER_CV APPELEE")
    print(f"Méthode: {request.method}")
    print(f"Fichiers: {request.FILES}")
    print("="*50)
    
    if request.method == 'POST':
        if not request.FILES.get('file'):
            print("❌ Aucun fichier trouvé")
            return JsonResponse({'error': 'Aucun fichier fourni'}, status=400)
        
        file = request.FILES['file']
        filename = file.name
        print(f"📁 Fichier reçu: {filename}")
        print(f"📊 Taille: {file.size} octets")
        
        try:
            if filename.endswith('.json'):
                # Lire le fichier JSON
                print("📄 Traitement fichier JSON...")
                data = json.load(file)
                
                # Créer le CV
                cv = CVGenere.objects.create(
                    utilisateur=request.user,
                    titre=data.get('titre', filename),
                    donnees_cv=data,
                    est_public=False
                )
                
                print(f"✅ CV JSON créé avec ID: {cv.id}")
                return JsonResponse({'success': True, 'id': cv.id})
                
            elif filename.endswith('.pdf'):
                print("📄 Traitement fichier PDF...")
                # Créer un CV minimal avec le fichier PDF
                cv = CVGenere.objects.create(
                    utilisateur=request.user,
                    titre=filename.replace('.pdf', ''),
                    donnees_cv={'prenom': filename.replace('.pdf', ''), 'nom': '', 'fichier_original': filename},
                    est_public=False
                )
                
                # Sauvegarder le fichier PDF
                cv.fichier_pdf.save(filename, ContentFile(file.read()))
                
                print(f"✅ CV PDF créé avec ID: {cv.id}")
                return JsonResponse({'success': True, 'id': cv.id})
                
            elif filename.endswith('.docx'):
                print("📄 Traitement fichier DOCX...")
                # Créer un CV minimal avec le fichier DOCX
                cv = CVGenere.objects.create(
                    utilisateur=request.user,
                    titre=filename.replace('.docx', ''),
                    donnees_cv={'prenom': filename.replace('.docx', ''), 'nom': '', 'fichier_original': filename},
                    est_public=False
                )
                
                # Sauvegarder le fichier DOCX
                cv.fichier_docx.save(filename, ContentFile(file.read()))
                
                print(f"✅ CV DOCX créé avec ID: {cv.id}")
                return JsonResponse({'success': True, 'id': cv.id})
                
            else:
                print(f"❌ Format non supporté: {filename}")
                return JsonResponse({'error': 'Format non supporté. Utilisez .pdf, .docx ou .json'}, status=400)
                
        except json.JSONDecodeError as e:
            print(f"❌ Erreur JSON: {e}")
            return JsonResponse({'error': 'Fichier JSON invalide'}, status=400)
        except Exception as e:
            print(f"❌ Erreur: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Méthode non autorisée'}, status=405)

@login_required
def modifier_cv(request, cv_id):
    """Modifier un CV - Version simple"""
    cv = get_object_or_404(CVGenere, id=cv_id, utilisateur=request.user)
    
    if request.method == 'POST':
        # Mettre à jour le titre
        cv.titre = request.POST.get('titre', cv.titre)
        cv.save()
        messages.success(request, "CV modifié avec succès")
        return redirect('myAppli:mes_cvs')
    
    return render(request, 'myAppli/outils_emploi/modifier_cv.html', {'cv': cv})

@login_required
def publier_offre_emploi(request):
    """
    Vue pour publier une offre d'emploi
    """
    if request.method == 'POST':
        # Vérifier que l'utilisateur est un recruteur
        if not hasattr(request.user, 'particulier') or not hasattr(request.user.particulier, 'recruteur'):
            messages.error(request, "Vous devez être recruteur pour publier une offre")
            return redirect('myAppli:home')
        
        recruteur = request.user.particulier.recruteur
        
        try:
            # Traitement des compétences
            competences_requises = []
            if request.POST.get('competences_requises'):
                competences_requises = [c.strip() for c in request.POST.get('competences_requises').split(',') if c.strip()]
            
            competences_souhaitees = []
            if request.POST.get('competences_souhaitees'):
                competences_souhaitees = [c.strip() for c in request.POST.get('competences_souhaitees').split(',') if c.strip()]
            
            # Création de l'offre
            offre = OffreEmploi.objects.create(
                recruteur=recruteur,
                titre=request.POST.get('titre'),
                description=request.POST.get('description', ''),
                missions=request.POST.get('missions', ''),
                profil_recherche=request.POST.get('profil_recherche', ''),
                localisation=request.POST.get('localisation'),
                type_contrat=request.POST.get('type_contrat'),
                teletravail='TOTAL' if request.POST.get('teletravail') else 'NON',
                niveau_experience=request.POST.get('niveau_experience'),
                annees_experience_min=request.POST.get('annees_experience_min') or 0,
                niveau_etude_requis=request.POST.get('niveau_etude_requis', ''),
                competences_requises=competences_requises,
                competences_souhaitees=competences_souhaitees,
                salaire_min=request.POST.get('salaire_min') or None,
                salaire_max=request.POST.get('salaire_max') or None,
                salaire_affiche=request.POST.get('salaire_affiche', ''),
                date_limite=request.POST.get('date_limite') or None,
                statut='PUBLIEE',
                est_active=True,
                source='MANUEL'
            )
            
            messages.success(request, f"L'offre '{offre.titre}' a été publiée avec succès !")
            logger.info(f"Offre publiée par {recruteur.organisation} : {offre.titre}")
            
        except Exception as e:
            messages.error(request, f"Erreur lors de la publication : {str(e)}")
            logger.error(f"Erreur publication offre : {e}")
        
        return redirect('myAppli:dashboard_recruteur')
    
    return redirect('myAppli:dashboard_recruteur')

def generer_lettre_motivation(request):
    """
    Générateur de lettre de motivation personnalisable
    """
    # Récupérer les données du candidat si connecté
    user_data = {}
    if request.user.is_authenticated and hasattr(request.user, 'particulier'):
        particulier = request.user.particulier
        if hasattr(particulier, 'candidat'):
            candidat = particulier.candidat
            user_data = {
                'nom': particulier.nom,
                'prenom': particulier.prenom,
                'email': particulier.email,
                'telephone': str(particulier.telephone),
                'adresse': particulier.adresse,
                'ville': particulier.ville,
                'competences': candidat.competences,
                'niveauEtude': candidat.niveauEtude,
                'experiences': candidat.anneesExperiences,
            }
    
    # Récupérer les offres d'emploi pour le select
    offres = OffreEmploi.objects.filter(statut='PUBLIEE', est_active=True).order_by('-date_publication')[:20]
    
    context = {
        'user_data': user_data,
        'offres': offres,
    }
    return render(request, 'myAppli/outils_emploi/generer_lettre_motivation.html', context)

@require_http_methods(["POST"])
def generer_lettre_pdf(request):
    """
    Génère le PDF de la lettre de motivation
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfgen import canvas
    import io
    from django.http import HttpResponse
    
    # Récupérer les données du formulaire
    data = request.POST
    prenom = data.get('prenom', '')
    nom = data.get('nom', '')
    email = data.get('email', '')
    telephone = data.get('telephone', '')
    adresse = data.get('adresse', '')
    ville = data.get('ville', '')
    
    entreprise_nom = data.get('entreprise_nom', '')
    poste = data.get('poste', '')
    offre_reference = data.get('offre_reference', '')
    
    # Sections personnalisées
    presentation = data.get('presentation', '')
    competences = data.get('competences', '')
    motivations = data.get('motivations', '')
    disponibilite = data.get('disponibilite', '')
    
    # Style choisi
    style = data.get('style', 'classique')
    
    # Créer le PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Styles personnalisés
    style_title = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1ed760'),
        alignment=TA_CENTER,
        spaceAfter=30,
    )
    
    style_subtitle = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceAfter=20,
    )
    
    style_normal = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
    )
    
    style_signature = ParagraphStyle(
        'Signature',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_LEFT,
        spaceBefore=30,
    )
    
    # En-tête selon le style
    if style == 'moderne':
        # En-tête moderne avec cadre
        header_data = [[
            Paragraph(f"<b>{prenom} {nom}</b>", style_title),
        ]]
        header_table = Table(header_data, colWidths=[doc.width])
        header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#1ed760')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 15),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 20))
        
        # Coordonnées
        contact_text = f"{adresse}<br/>{ville}<br/>{telephone}<br/>{email}"
        story.append(Paragraph(contact_text, styles['Normal']))
        
    else:  # classique
        # En-tête classique
        story.append(Paragraph(f"{prenom} {nom}", style_title))
        story.append(Paragraph(f"{adresse}", styles['Normal']))
        story.append(Paragraph(f"{ville}", styles['Normal']))
        story.append(Paragraph(f"Tél: {telephone} | Email: {email}", styles['Normal']))
    
    story.append(Spacer(1, 30))
    
    # Date
    from datetime import datetime
    date_obj = timezone.now()()
    story.append(Paragraph(f"Le {date_obj.strftime('%d/%m/%Y')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Destinataire
    story.append(Paragraph(f"À l'attention du recruteur", styles['Normal']))
    story.append(Paragraph(f"{entreprise_nom}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Objet
    story.append(Paragraph(f"<b>Objet : Candidature pour le poste de {poste}</b>", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Corps de la lettre
    story.append(Paragraph("Madame, Monsieur,", styles['Normal']))
    story.append(Spacer(1, 10))
    
    # Introduction
    intro = f"Actuellement à la recherche d'un nouveau défi professionnel, j'ai l'honneur de vous présenter ma candidature pour le poste de <b>{poste}</b> au sein de votre structure."
    if offre_reference:
        intro += f" Suite à votre offre d'emploi référencée <b>{offre_reference}</b>, je me permets de vous adresser ma candidature."
    story.append(Paragraph(intro, style_normal))
    story.append(Spacer(1, 10))
    
    # Présentation
    if presentation:
        story.append(Paragraph(presentation, style_normal))
        story.append(Spacer(1, 10))
    else:
        story.append(Paragraph(f"Titulaire d'un {user_data.get('niveauEtude', 'diplôme')}, je dispose d'une expérience de {user_data.get('experiences', 0)} ans dans le domaine.", style_normal))
        story.append(Spacer(1, 10))
    
    # Compétences
    if competences:
        story.append(Paragraph("<b>Mes compétences clés :</b>", styles['Normal']))
        story.append(Paragraph(competences, style_normal))
    else:
        story.append(Paragraph("<b>Mes atouts :</b>", styles['Normal']))
        story.append(Paragraph(f"• {user_data.get('competences', 'Rigueur, autonomie et esprit d\'équipe')}", style_normal))
    story.append(Spacer(1, 10))
    
    # Motivations
    if motivations:
        story.append(Paragraph("<b>Mes motivations :</b>", styles['Normal']))
        story.append(Paragraph(motivations, style_normal))
    else:
        story.append(Paragraph("<b>Pourquoi me rejoindre ?</b>", styles['Normal']))
        story.append(Paragraph("Je suis convaincu que mon profil correspond aux valeurs et aux besoins de votre entreprise. Dynamique et passionné, je saurai m'investir pleinement dans les missions qui me seront confiées.", style_normal))
    story.append(Spacer(1, 10))
    
    # Disponibilité
    if disponibilite:
        story.append(Paragraph(f"Je suis disponible à compter du {disponibilite}.", styles['Normal']))
    else:
        story.append(Paragraph("Je me tiens à votre disposition pour un entretien à votre convenance.", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Formule de politesse
    story.append(Paragraph("Je vous remercie de l'attention que vous porterez à ma candidature et vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.", style_normal))
    story.append(Spacer(1, 30))
    
    # Signature
    story.append(Paragraph(f"{prenom} {nom}", style_signature))
    
    # Générer le PDF
    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()
    
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="lettre_motivation_{nom}_{prenom}.pdf"'
    return response

def telecharger_modele_lettre(request):
    """
    Télécharge un modèle de lettre de motivation au format DOCX
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    buffer = io.BytesIO()
    doc = Document()
    
    # Style du document
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # En-tête
    title = doc.add_heading('LETTRE DE MOTIVATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Coordonnées
    doc.add_paragraph('[Votre Nom et Prénom]')
    doc.add_paragraph('[Votre Adresse]')
    doc.add_paragraph('[Code Postal, Ville]')
    doc.add_paragraph('[Téléphone]')
    doc.add_paragraph('[Email]')
    
    doc.add_paragraph()
    doc.add_paragraph(f"Le [Date]")
    doc.add_paragraph()
    
    # Destinataire
    doc.add_paragraph("À l'attention du service recrutement")
    doc.add_paragraph("[Nom de l'entreprise]")
    doc.add_paragraph("[Adresse de l'entreprise]")
    
    doc.add_paragraph()
    doc.add_paragraph("Objet : Candidature pour le poste de [Intitulé du poste]")
    doc.add_paragraph()
    
    # Corps
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()
    
    p = doc.add_paragraph("Actuellement à la recherche d'un nouveau défi professionnel, j'ai l'honneur de vous présenter ma candidature pour le poste de [Intitulé du poste] au sein de votre structure.")
    doc.add_paragraph()
    
    doc.add_paragraph("[Décrivez votre parcours et vos compétences]")
    doc.add_paragraph()
    
    doc.add_paragraph("Je reste à votre disposition pour un entretien afin de vous exposer plus en détail ma motivation et mes compétences.")
    doc.add_paragraph()
    
    doc.add_paragraph("Je vous remercie de l'attention que vous porterez à ma candidature et vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.")
    doc.add_paragraph()
    
    doc.add_paragraph("[Signature]")
    
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="modele_lettre_motivation.docx"'
    return response

def apercu_style_lettre(request, style):
    """Génère un aperçu HTML d'un style de lettre"""
    # Données d'exemple pour l'aperçu
    contexte = {
        'style': style,
        'prenom': 'Juste',
        'nom': 'KABORE',
        'email': 'juste.kabore@email.com',
        'telephone': '01 23 45 67',
        'adresse': '123 rue de l\'Exemple',
        'ville': 'Bobo',
        'entreprise_nom': 'Entreprise ABC',
        'poste': 'Développeur Web',
        'date': datetime.now().strftime('%d/%m/%Y')
    }
    
    return render(request, f'lettre/aperçu_{style}.html', contexte)

def gestion_entreprise(request):
    # Liste des outils en développement
    tools_in_development = [
        'facture', 'offre_cout', 'contrat', 'fiche_poste', 'paie', 
        'rapport', 'pv', 'pitch', 'communique', 'plan_com', 
        'profil_ent', 'note_meth', 'biblio_doc'
    ]
    
    # Statistiques pour le hero
    stats = {
        'total_outils': 21,
        'categories': 5,
        'ia_assistee': 6
    }
    
    # Données de démonstration pour le tableau de bord
    dashboard_data = {
        'kpis': {
            'chiffre_affaire_mois': 24500,
            'objectif_mensuel': 30000,
            'taux_realisation': 82,
            'nb_candidats_mois': 12,
            'nb_entretiens_realises': 8,
            'projets_en_cours': 3,
            'taux_activite': 75
        },
        'alertes': [
            {'type': 'warning', 'message': 'Déclaration URSSAF dans 5 jours'},
            {'type': 'info', 'message': 'Entretien annuel à planifier pour 3 employés'},
            {'type': 'success', 'message': 'Votre proposition pour le marché SONABHY a été retenue !'}
        ]
    }
    
    # Date du jour pour les calculs
    today = timezone.now().date()
    
    context = {
        'today': today,
        'tools_in_development': tools_in_development,
        'stats': stats,
        'dashboard': dashboard_data,
        'annee_courante': today.year,
    }
    
    return render(request, 'myAppli/gestion/gestion_entreprise.html', context)

@login_required
def soumettre_candidature(request, offre_id):
    """
    Page de préparation du dossier de candidature
    """
    if not hasattr(request.user, 'particulier'):
        return redirect('myAppli:home')

    candidat = request.user.particulier.candidat

    try:
        offre = OffreEmploi.objects.get(id=offre_id, est_active=True)
    except OffreEmploi.DoesNotExist:
        messages.error(request, "Offre introuvable")
        return redirect('myAppli:dashboard_candidat')

    # Vérifier si candidature déjà soumise
    dossier_existant = DossierCandidature.objects.filter(
        candidat=candidat,
        offre=offre
    ).first()

    if dossier_existant and dossier_existant.statut == DossierCandidature.SOUMIS:
        messages.warning(request, "Vous avez déjà soumis une candidature pour cette offre.")
        return redirect('myAppli:detail_dossier_candidature', dossier_id=dossier_existant.id)

    # Documents exigés par l'offre
    documents_exiges = ['CV', 'Lettre de motivation']
    try:
        ct = ContentType.objects.get_for_model(OffreEmploi)
        analyse = AnalyseDocument.objects.get(content_type=ct, object_id=offre.id)
        documents_exiges = analyse.elements_offre.documents_exiges
    except Exception:
        pass

    # CVs disponibles du candidat
    cvs_generes = CVGenere.objects.filter(
        utilisateur=candidat.particulier.user,
        fichier_pdf__isnull=False
    ).order_by('-date_generation')

    # Score de compatibilité
    score = 0
    try:
        reco = RecommandationEmploi.objects.get(candidat=candidat, offre=offre)
        score = round(reco.score_global * 100, 1)
    except RecommandationEmploi.DoesNotExist:
        pass

    lettres_generees = LettreMotivationGeneree.objects.filter(
        utilisateur=candidat.particulier.user
    ).order_by('-date_generation')[:5]

    return render(request, 'myAppli/candidature/soumettre_candidature.html', {
        'candidat': candidat,
        'offre': offre,
        'documents_exiges': documents_exiges,
        'cvs_generes': cvs_generes,
        'dossier_existant': dossier_existant,
        'score': score,
        'lettres_generees': lettres_generees,
    })

@login_required
def analyser_cv_offre(request, offre_id):
    """
    Analyse IA du CV par rapport à l'offre via AJAX
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)

    candidat = request.user.particulier.candidat
    offre = OffreEmploi.objects.get(id=offre_id)
    cv_genere_id = request.POST.get('cv_genere_id')

    # Récupérer le texte du CV
    cv_texte = ""
    if cv_genere_id:
        try:
            cv = CVGenere.objects.get(id=cv_genere_id, utilisateur=request.user)
            import json
            d = cv.donnees_cv
            cv_texte = f"""
            Nom: {d.get('prenom', '')} {d.get('nom', '')}
            Compétences: {', '.join([c.get('nom', c) if isinstance(c, dict) else c for c in d.get('competences', [])])}
            Expériences: {len(d.get('experiences', []))} expériences
            Formations: {', '.join([f.get('diplome', '') if isinstance(f, dict) else f for f in d.get('formations', [])])}
            Langues: {', '.join([l.get('nom', l) if isinstance(l, dict) else l for l in d.get('langues', [])])}
            """
        except CVGenere.DoesNotExist:
            pass
    else:
        from analyse_ia.analyser_pdfs import extraire_texte_cv
        cv_texte = extraire_texte_cv(candidat.particulier_id) or candidat.competences

    # Analyse via Groq
    try:
        from analyse_ia.ia_client import IAClient
        ia = IAClient()

        prompt = f"""
        Tu es un recruteur expert. Analyse la compatibilité entre ce CV et cette offre.

        OFFRE:
        Titre: {offre.titre}
        Compétences requises: {', '.join(offre.competences_requises)}
        Niveau requis: {offre.niveau_etude_requis}
        Expérience requise: {offre.annees_experience_min} ans minimum
        Description: {offre.description[:300]}

        CV DU CANDIDAT:
        {cv_texte[:500]}

        Réponds en JSON :
        {{
            "score": 75,
            "points_forts": ["Point 1", "Point 2"],
            "points_faibles": ["Point 1", "Point 2"],
            "conseil": "Conseil pour améliorer la candidature"
        }}
        """

        response = ia.groq.chat.completions.create(
            model=ia.groq_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=500,
        )

        import json
        texte = response.choices[0].message.content.strip()
        if '```json' in texte:
            texte = texte.split('```json')[1].split('```')[0]
        elif '```' in texte:
            texte = texte.split('```')[1].split('```')[0]

        data = json.loads(texte.strip())
        return JsonResponse({
            'success': True,
            'score': data.get('score', 0),
            'points_forts': data.get('points_forts', []),
            'points_faibles': data.get('points_faibles', []),
            'conseil': data.get('conseil', ''),
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def generer_lettre(request, offre_id):
    """
    Génère une lettre de motivation via IA
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Méthode non autorisée'}, status=405)

    candidat = request.user.particulier.candidat
    offre = OffreEmploi.objects.get(id=offre_id)

    try:
        from analyse_ia.ia_client import IAClient
        ia = IAClient()

        prompt = f"""
        Tu es un expert en rédaction de lettres de motivation.
        Génère une lettre de motivation professionnelle en français.

        CANDIDAT:
        Nom: {candidat.prenom} {candidat.nom}
        Compétences: {candidat.competences}
        Niveau d'étude: {candidat.niveauEtude}
        Années d'expérience: {candidat.anneesExperiences}
        Secteur: {candidat.secteur_recherche}

        OFFRE:
        Titre: {offre.titre}
        Entreprise: {offre.entreprise_nom or 'l\'entreprise'}
        Description: {offre.description[:300]}
        Compétences requises: {', '.join(offre.competences_requises)}

        Génère une lettre de motivation complète et professionnelle.
        La lettre doit être personnalisée, convaincante et adaptée au poste.
        Format: texte simple sans markdown, avec paragraphes séparés par des sauts de ligne.
        """

        response = ia.groq.chat.completions.create(
            model=ia.groq_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1000,
        )

        lettre_texte = response.choices[0].message.content.strip()
        
        # Sauvegarder dans LettreMotivationGeneree
        lettre_obj = LettreMotivationGeneree.objects.create(
            utilisateur=request.user,
            candidat=candidat,
            offre=offre,
            titre=f"Lettre — {offre.titre} — {candidat.prenom} {candidat.nom}",
            contenu=lettre_texte,
        )

        return JsonResponse({
            'success': True, 
            'lettre': lettre_texte,
            'lettre_id': lettre_obj.id,
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@login_required
def enregistrer_dossier(request, offre_id):
    """
    Enregistre le dossier de candidature
    """
    if request.method != 'POST':
        return redirect('myAppli:soumettre_candidature', offre_id=offre_id)

    candidat = request.user.particulier.candidat
    offre = OffreEmploi.objects.get(id=offre_id)

    # Récupérer ou créer le dossier
    dossier, created = DossierCandidature.objects.get_or_create(
        candidat=candidat,
        offre=offre
    )

    # CV
    cv_genere_id = request.POST.get('cv_genere_id')
    if cv_genere_id:
        try:
            dossier.cv_genere = CVGenere.objects.get(
                id=cv_genere_id,
                utilisateur=request.user
            )
        except CVGenere.DoesNotExist:
            pass
    elif 'cv' in request.FILES:
        dossier.cv = request.FILES['cv']

    # Lettre de motivation
    lettre_id = request.POST.get('lettre_id')
    lettre_generee_texte = request.POST.get('lettre_generee', '')

    if lettre_id:
        try:
            from myAppli.models import LettreMotivationGeneree
            lettre_obj = LettreMotivationGeneree.objects.get(
                id=lettre_id,
                utilisateur=request.user
            )
            dossier.lettre_generee_obj = lettre_obj
            dossier.lettre_generee = lettre_generee_texte
        except LettreMotivationGeneree.DoesNotExist:
            pass
    elif lettre_generee_texte:
        dossier.lettre_generee = lettre_generee_texte
    elif 'lettre_motivation' in request.FILES:
        dossier.lettre_motivation = request.FILES['lettre_motivation']
    
    # Autres documents
    if 'autres_documents' in request.FILES:
        dossier.autres_documents = request.FILES['autres_documents']

    # Message
    dossier.message_candidat = request.POST.get('message_candidat', '')

    action = request.POST.get('action', 'sauvegarder')

    if action in ['soumettre', 'telecharger']:
        erreurs = []

        # Vérifier CV
        a_cv = (
            cv_genere_id or
            'cv' in request.FILES or
            candidat.cv or
            CVGenere.objects.filter(
                utilisateur=request.user,
                fichier_pdf__isnull=False
            ).exists()
        )
        if not a_cv:
            erreurs.append('CV manquant')

        # Vérifier lettre
        a_lettre = (
            request.POST.get('lettre_generee', '').strip() or
            'lettre_motivation' in request.FILES
        )
        if not a_lettre:
            erreurs.append('Lettre de motivation manquante')

        if erreurs:
            for err in erreurs:
                messages.error(request, f"❌ {err}")
            return redirect('myAppli:soumettre_candidature', offre_id=offre_id)

    dossier.save()

    if action == 'soumettre':
        # ✅ Calculer les scores IA avec Ollama
        print(f"🤖 Calcul des scores IA pour le dossier #{dossier.id}...")
        succes_ia = mettre_a_jour_scores_dossier(dossier.id, force=False)
        
        if succes_ia:
            print(f"✅ Scores IA calculés avec succès !")
        else:
            print(f"⚠️ Échec du calcul IA, scores par défaut conservés")
        
        dossier.refresh_from_db()
        
        dossier.statut = DossierCandidature.SOUMIS
        dossier.date_soumission = timezone.now()
        dossier.save()

        # Incrémenter les compteurs
        offre.incrementer_candidatures()
        candidat.nb_candidatures_envoyees += 1
        candidat.save(update_fields=['nb_candidatures_envoyees'])

        # Envoyer email au recruteur
        _envoyer_email_recruteur(dossier, offre, candidat)

        messages.success(request, "🎉 Candidature soumise avec succès !")
        return redirect('myAppli:detail_dossier_candidature', dossier_id=dossier.id)

    elif action == 'telecharger':
        dossier.save()
        return _telecharger_dossier_zip(dossier, offre, candidat)

    else:
        messages.success(request, "Dossier sauvegardé !")
        return redirect('myAppli:soumettre_candidature', offre_id=offre_id)

def _envoyer_email_recruteur(dossier, offre, candidat):
    """Envoie un email au recruteur"""
    try:
        email_recruteur = (
            offre.recruteur.email if offre.recruteur
            else offre.entreprise_email
        )

        if not email_recruteur:
            print("⚠️ Pas d'email recruteur trouvé")
            return

        sujet = f"Candidature — {offre.titre} — {candidat.prenom} {candidat.nom}"

        corps = f"""
Bonjour,

{candidat.prenom} {candidat.nom} a soumis une candidature pour le poste de {offre.titre}.

Informations du candidat :
- Niveau d'étude : {candidat.niveauEtude}
- Années d'expérience : {candidat.anneesExperiences} ans
- Compétences : {candidat.competences}
- Secteur : {candidat.secteur_recherche}
- Score de compatibilité : {round(dossier.score_compatibilite * 100, 1)}%

Message du candidat :
{dossier.message_candidat or 'Aucun message'}

---
Candidature reçue via FASOIA
        """

        email = EmailMessage(
            subject=sujet,
            body=corps,
            from_email='noreply@fasoia.com',
            to=[email_recruteur],
        )

        # Joindre le CV
        if dossier.cv:
            email.attach_file(dossier.cv.path)
        elif dossier.cv_genere and dossier.cv_genere.fichier_pdf:
            email.attach_file(dossier.cv_genere.fichier_pdf.path)

        email.send()
        print(f"✅ Email envoyé à {email_recruteur}")

    except Exception as e:
        print(f"❌ Erreur envoi email: {e}")

def _telecharger_dossier_zip(dossier, offre, candidat):
    """Génère et retourne un ZIP du dossier"""
    import io as io_module

    buffer = io_module.BytesIO()

    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

        # CV
        if dossier.cv and dossier.cv.path:
            zip_file.write(dossier.cv.path, f"CV_{candidat.nom}_{candidat.prenom}.pdf")
        elif dossier.cv_genere and dossier.cv_genere.fichier_pdf:
            zip_file.write(
                dossier.cv_genere.fichier_pdf.path,
                f"CV_{candidat.nom}_{candidat.prenom}.pdf"
            )

        # Lettre de motivation
        if dossier.lettre_motivation:
            zip_file.write(
                dossier.lettre_motivation.path,
                f"Lettre_motivation_{candidat.nom}_{candidat.prenom}.pdf"
            )
        elif dossier.lettre_generee:
            # Convertir la lettre texte en fichier
            zip_file.writestr(
                f"Lettre_motivation_{candidat.nom}_{candidat.prenom}.txt",
                dossier.lettre_generee
            )

        # Autres documents
        if dossier.autres_documents:
            zip_file.write(
                dossier.autres_documents.path,
                f"Documents_{candidat.nom}_{candidat.prenom}.pdf"
            )

    buffer.seek(0)
    nom_fichier = f"Candidature_{candidat.nom}_{candidat.prenom}_{offre.titre}.zip"

    response = HttpResponse(buffer.read(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{nom_fichier}"'
    return response

@login_required
def detail_dossier_candidature(request, dossier_id):
    """
    Détail d'un dossier de candidature
    """
    candidat = request.user.particulier.candidat

    try:
        dossier = DossierCandidature.objects.get(
            id=dossier_id,
            candidat=candidat
        )
    except DossierCandidature.DoesNotExist:
        messages.error(request, "Dossier introuvable")
        return redirect('myAppli:dashboard_candidat')

    return render(request, 'myAppli/candidature/detail_dossier.html', {
        'dossier': dossier,
        'offre': dossier.offre,
        'candidat': candidat,
    })