import os
import base64
from io import BytesIO
from django.conf import settings
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from django.core.files.base import ContentFile


class GenerateurCVPublic:
    """
    Service de génération de CV accessible à tous.
    Gère les 6 styles : executive, geometric, diagonal, dark, swiss, splitcard
    """


    def __init__(self):
        self.templates_dir = 'cv/'

    # ──────────────────────────────────────────────────────────────────────────
    # PRÉPARATION DES DONNÉES
    # ──────────────────────────────────────────────────────────────────────────

    def preparer_donnees(self, form_data, files=None):
        """
        Prépare les données du formulaire pour le template Django.
        Accepte request.POST et optionnellement request.FILES pour la photo.
        """

        # ── Expériences ───────────────────────────────────────────────────────
        experiences = []
        if form_data.get('experiences'):
            for ligne in form_data['experiences'].split('\n'):
                if ligne.strip():
                    parts = [p.strip() for p in ligne.split('|')]
                    if len(parts) >= 3:
                        experiences.append({
                            'titre':       parts[0],
                            'entreprise':  parts[1] if len(parts) > 1 else '',
                            'date_debut':  parts[2] if len(parts) > 2 else '',
                            'date_fin':    parts[3] if len(parts) > 3 else 'Présent',
                            'description': parts[4] if len(parts) > 4 else '',
                        })

        # ── Formations ────────────────────────────────────────────────────────
        formations = []
        if form_data.get('formations'):
            for ligne in form_data['formations'].split('\n'):
                if ligne.strip():
                    parts = [p.strip() for p in ligne.split('|')]
                    if len(parts) >= 3:
                        formations.append({
                            'diplome':        parts[0],
                            'etablissement':  parts[1] if len(parts) > 1 else '',
                            'annee':          parts[2] if len(parts) > 2 else '',
                        })

        # Après la ligne 54 (après le append)
        print(f"🔍 DEBUG formations trouvées : {len(formations)}")
        for f in formations:
            print(f"   - {f}")

        # ── Listes simples ────────────────────────────────────────────────────
        competences = []
        if form_data.get('competences'):
            competences = [c.strip() for c in form_data['competences'].split(',') if c.strip()]

        langues = []
        if form_data.get('langues'):
            langues = [l.strip() for l in form_data['langues'].split(',') if l.strip()]

        centres_interet = []
        if form_data.get('centres_interet'):
            centres_interet = [ci.strip() for ci in form_data['centres_interet'].split(',') if ci.strip()]

        # ── Photo ─────────────────────────────────────────────────────────────
        photo_src = None
        if files and files.get('photo'):
            photo = files['photo']
            try:
                photo_bytes = photo.read()
                photo_b64   = base64.b64encode(photo_bytes).decode('utf-8')
                photo_src   = f"data:{photo.content_type};base64,{photo_b64}"
            except Exception as e:
                print(f"⚠️ Erreur lecture photo : {e}")
                photo_src = None

        return {
            'style':            form_data.get('style', 'executive'),
            'prenom':           form_data.get('prenom', '').strip(),
            'nom':              form_data.get('nom', '').strip(),
            'email':            form_data.get('email', '').strip(),
            'telephone':        form_data.get('telephone', '').strip(),
            'adresse':          form_data.get('adresse', '').strip(),
            'ville':            form_data.get('ville', '').strip(),
            'pays':             form_data.get('pays', '').strip(),
            'code_postal':      form_data.get('code_postal', '').strip(),
            'permis':           form_data.get('permis', '').strip(),
            'resume':           form_data.get('resume', '').strip(),
            'competences':      competences,
            'langues':          langues,
            'centres_interet':  centres_interet,
            'experiences':      experiences,
            'formations':       formations,
            'photo_src':        photo_src,   # ← base64 ou None
            'date_generation':  datetime.now().strftime('%d/%m/%Y'),
            'DEBUG_formations_count': len(formations), 
        }

    # ──────────────────────────────────────────────────────────────────────────
    # GÉNÉRATION HTML
    # ──────────────────────────────────────────────────────────────────────────

    def generer_html(self, donnees):
        """Génère le HTML à partir du template Django."""
        style         = donnees.get('style', 'executive')
        template_name = f"{self.templates_dir}{style}.html"

        try:
            return render_to_string(template_name, donnees)
        except Exception as e:
            print(f"❌ Erreur template {template_name} : {e}")
            return self._html_secours(donnees)

    def generer_html_avec_donnees(self, style, donnees):
        """Génère le HTML à partir du template avec les données fournies."""
        donnees['style'] = style
        template_name   = f"{self.templates_dir}{style}.html"
        print(f"🔍 Template recherché : {template_name}")

        try:
            html = render_to_string(template_name, donnees)
            print("✅ Template trouvé, HTML généré")
            return html
        except Exception as e:
            print(f"❌ Erreur : {e}")
            return self._html_secours(donnees)

    # ──────────────────────────────────────────────────────────────────────────
    # GÉNÉRATION PDF / DOCX
    # ──────────────────────────────────────────────────────────────────────────

    def generer_pdf(self, donnees):
        """Génère un PDF à partir du HTML."""
        html_string = self.generer_html(donnees)
        return HTML(string=html_string, base_url=str(settings.BASE_DIR)).write_pdf()

    def generer_docx(self, donnees):
        """Génère un DOCX avec python-docx."""
        doc = Document()

        doc.add_heading(f"{donnees['prenom']} {donnees['nom']}", 0)

        if donnees.get('email') or donnees.get('telephone'):
            contact = doc.add_paragraph()
            if donnees.get('email'):
                contact.add_run(f"Email: {donnees['email']}\n")
            if donnees.get('telephone'):
                contact.add_run(f"Tél: {donnees['telephone']}")

        if donnees.get('resume'):
            doc.add_heading('PROFIL', level=1)
            doc.add_paragraph(donnees['resume'])

        if donnees.get('experiences'):
            doc.add_heading('EXPÉRIENCES PROFESSIONNELLES', level=1)
            for exp in donnees['experiences']:
                p = doc.add_paragraph()
                p.add_run(f"{exp['titre']} - {exp['entreprise']}").bold = True
                p.add_run(f" ({exp['date_debut']} - {exp.get('date_fin', 'Présent')})")
                if exp.get('description'):
                    doc.add_paragraph(exp['description'], style='List Bullet')

        if donnees.get('formations'):
            doc.add_heading('FORMATION', level=1)
            for formation in donnees['formations']:
                p = doc.add_paragraph()
                p.add_run(f"{formation['diplome']} - {formation['etablissement']}").bold = True
                p.add_run(f" ({formation['annee']})")

        if donnees.get('competences'):
            doc.add_heading('COMPÉTENCES', level=1)
            doc.add_paragraph(', '.join(donnees['competences']))

        if donnees.get('langues'):
            doc.add_heading('LANGUES', level=1)
            doc.add_paragraph(', '.join(donnees['langues']))

        if donnees.get('centres_interet'):
            doc.add_heading("CENTRES D'INTÉRÊT", level=1)
            doc.add_paragraph(', '.join(donnees['centres_interet']))

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer.getvalue()

    # ──────────────────────────────────────────────────────────────────────────
    # GÉNÉRATION BUFFER (utilisé par la vue)
    # ──────────────────────────────────────────────────────────────────────────

    def generer_cv_buffer(self, donnees, format_export, style='executive'):
        """
        Génère le CV et retourne (buffer, nom_fichier).
        Utilisé pour la sauvegarde en base avant téléchargement.
        """
        donnees['style'] = style
        buffer           = BytesIO()
        nom_base         = (
            f"CV_{donnees['prenom']}_{donnees['nom']}_"
            f"{donnees.get('date_generation', datetime.now().strftime('%Y%m%d_%H%M%S'))}"
        ).replace(' ', '_')

        if format_export == 'pdf':
            buffer.write(self.generer_pdf(donnees))
            extension = 'pdf'
        elif format_export == 'docx':
            buffer.write(self.generer_docx(donnees))
            extension = 'docx'
        else:
            raise ValueError(f"Format non supporté : {format_export}")

        buffer.seek(0)
        return buffer, f"{nom_base}.{extension}"

    def generer_cv(self, donnees, format='pdf'):
        """Génère un CV et retourne une HttpResponse en téléchargement."""
        if format == 'pdf':
            contenu      = self.generer_pdf(donnees)
            content_type = 'application/pdf'
            extension    = 'pdf'
        elif format == 'docx':
            contenu      = self.generer_docx(donnees)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            extension    = 'docx'
        else:
            raise ValueError(f"Format non supporté : {format}")

        nom_fichier = (
            f"CV_{donnees['prenom']}_{donnees['nom']}_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        ).replace(' ', '_')

        response = HttpResponse(contenu, content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.{extension}"'
        return response

    # ──────────────────────────────────────────────────────────────────────────
    # APERÇU
    # ──────────────────────────────────────────────────────────────────────────

    def generer_apercu(self, style):
        """Génère un aperçu HTML du style avec données d'exemple."""
        donnees_exemple = {
            'style':           style,
            'prenom':          'Ramata',
            'nom':             'TALL',
            'email':           'ramata.tall@fasoia.com',
            'telephone':       '+226 09 09 09 56 78',
            'adresse':         '15 Rue de la Paix',
            'ville':           'Patte d\'Oie',
            'pays':            'BURKINA FASO',
            'code_postal':     '',
            'permis':          'B',
            'resume':          'Professionnelle dynamique avec 5 ans d\'expérience dans le développement web. Passionnée par les nouvelles technologies et le travail en équipe.',
            'photo_src':       None,
            'competences':     ['Python', 'Django', 'JavaScript', 'React', 'SQL', 'Git'],
            'langues':         ['Français (courant)', 'Anglais (professionnel)', 'Mooré (natif)'],
            'centres_interet': ['Lecture', 'Sport', 'Voyages', 'Photographie'],
            'experiences': [
                {
                    'titre':       'Développeuse Full Stack',
                    'entreprise':  'TechCorp',
                    'date_debut':  '2021',
                    'date_fin':    'Présent',
                    'description': 'Développement d\'applications web avec Django et React. Gestion d\'une équipe de 3 développeurs.',
                },
                {
                    'titre':       'Développeuse Junior',
                    'entreprise':  'StartupIA',
                    'date_debut':  '2019',
                    'date_fin':    '2021',
                    'description': 'Maintenance et évolution d\'applications existantes. Participation à la refonte du site principal.',
                },
            ],
            'formations': [
                {
                    'diplome':        'Master en Informatique',
                    'etablissement':  'Université Joseph KI-ZERBO',
                    'annee':          '2019',
                },
                {
                    'diplome':        'Licence en Mathématiques',
                    'etablissement':  'Université de Ouagadougou',
                    'annee':          '2017',
                },
            ],
            'date_generation': datetime.now().strftime('%d/%m/%Y'),
        }

        template_name = f"{self.templates_dir}{style}.html"
        try:
            return render_to_string(template_name, donnees_exemple)
        except Exception as e:
            print(f"❌ Erreur aperçu {style} : {e}")
            return self._html_secours(donnees_exemple)

    # ──────────────────────────────────────────────────────────────────────────
    # SAUVEGARDE EN BASE
    # ──────────────────────────────────────────────────────────────────────────

    def sauvegarder_cv(self, donnees, contenu, format_fichier, user):
        """Sauvegarde le CV en base de données."""
        from ..models import CVGenere, ModeleCV

        modele = None
        try:
            modele = ModeleCV.objects.get(categorie=donnees.get('style', 'executive'), est_actif=True)
            modele.nb_utilisations += 1
            modele.save(update_fields=['nb_utilisations'])
        except ModeleCV.DoesNotExist:
            pass

        cv = CVGenere.objects.create(
            utilisateur=user,
            modele=modele,
            titre=f"CV - {donnees['prenom']} {donnees['nom']}",
            donnees_cv=donnees,
        )

        nom_fichier = (
            f"CV_{donnees['prenom']}_{donnees['nom']}_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        ).replace(' ', '_')

        if format_fichier == 'pdf':
            cv.fichier_pdf.save(f"{nom_fichier}.pdf", ContentFile(contenu), save=True)
        elif format_fichier == 'docx':
            cv.fichier_docx.save(f"{nom_fichier}.docx", ContentFile(contenu), save=True)

        return cv

    # ──────────────────────────────────────────────────────────────────────────
    # HELPERS PRIVÉS
    # ──────────────────────────────────────────────────────────────────────────

    def _html_secours(self, donnees):
        """Template HTML minimaliste en cas d'erreur."""
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; padding: 40px; max-width: 800px; margin: 0 auto; }}
                h1 {{ color: #1852b4; }}
                h2 {{ color: #1852b4; border-bottom: 2px solid #1852b4; padding-bottom: 5px; margin-top: 30px; }}
                .contact {{ color: #607897; margin-bottom: 20px; }}
            </style>
        </head>
        <body>
            <h1>{donnees.get('prenom', '')} {donnees.get('nom', '')}</h1>
            <div class="contact">
                {donnees.get('email', '')}
                {'&nbsp;|&nbsp;' + donnees.get('telephone', '') if donnees.get('telephone') else ''}
            </div>
            {f"<p>{donnees.get('resume', '')}</p>" if donnees.get('resume') else ''}
            {"<h2>Compétences</h2><p>" + ', '.join(donnees.get('competences', [])) + "</p>" if donnees.get('competences') else ''}
        </body>
        </html>
        """
