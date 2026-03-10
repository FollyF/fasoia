# myAppli/services/generateur_document.py
import os
from docx import Document
from datetime import datetime
from django.conf import settings
import hashlib

class GenerateurDocument:
    """Service de génération de documents utilisé par les vues"""
    
    def __init__(self):
        self.output_dir = os.path.join(settings.MEDIA_ROOT, 'soumissions/documents/')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def collecter_donnees_entreprise(self, entreprise):
        """Collecte toutes les données de l'entreprise"""
        donnees = {
            '{{ENTREPRISE_RAISON_SOCIALE}}': entreprise.raisonSociale,
            '{{ENTREPRISE_DOMAINE}}': entreprise.domaineActive,
            '{{ENTREPRISE_COMPETENCES}}': entreprise.competencesCles,
            '{{ENTREPRISE_LOCALISATION}}': entreprise.localisation,
            '{{ENTREPRISE_TAILLE}}': str(entreprise.taille),
            '{{ENTREPRISE_DESCRIPTION}}': entreprise.description,
            '{{ENTREPRISE_SITE_WEB}}': entreprise.site_web,
            '{{ENTREPRISE_ANNEE_CREATION}}': str(entreprise.annee_creation or ''),
            '{{ENTREPRISE_CHIFFRE_AFFAIRES}}': str(entreprise.chiffre_affaires or ''),
            '{{ENTREPRISE_CAPITAL}}': str(entreprise.capital_social or ''),
            '{{ENTREPRISE_EXPERIENCE}}': str(entreprise.annees_experience),
            '{{ENTREPRISE_PROJETS}}': str(entreprise.nb_projets_realises),
            '{{ENTREPRISE_REFERENCES}}': entreprise.references,
            '{{CONTACT_NOM}}': entreprise.nom,
            '{{CONTACT_PRENOM}}': entreprise.prenom,
            '{{CONTACT_EMAIL}}': entreprise.email,
            '{{CONTACT_TELEPHONE}}': str(entreprise.telephone),
        }
        return donnees
    
    def collecter_donnees_opportunite(self, opportunite, opportunite_type):
        """Collecte les données de l'opportunité"""
        donnees = {
            '{{OPPORTUNITE_REFERENCE}}': getattr(opportunite, 'reference', ''),
            '{{OPPORTUNITE_TITRE}}': getattr(opportunite, 'titre', ''),
            '{{OPPORTUNITE_DESCRIPTION}}': getattr(opportunite, 'description', ''),
            '{{OPPORTUNITE_SECTEUR}}': getattr(opportunite, 'secteur', ''),
            '{{OPPORTUNITE_DATE_LIMITE}}': getattr(opportunite, 'dateLimite', '').strftime('%d/%m/%Y') if hasattr(opportunite, 'dateLimite') and opportunite.dateLimite else '',
        }
        
        # Données spécifiques selon le type
        if opportunite_type == 'Offre_uemoa':
            donnees.update({
                '{{OFFRE_DOWNLOAD_URL}}': getattr(opportunite, 'download_url', ''),
            })
        else:  # Ami_uemoa
            donnees.update({
                '{{AMI_OBJET}}': getattr(opportunite, 'objet', ''),
                '{{AMI_DOCUMENTS_REQUIS}}': getattr(opportunite, 'documentsRequis', ''),
            })
        
        return donnees
    
    def generer(self, modele, entreprise, opportunite, opportunite_type, donnees_supp=None):
        """
        Génère un document à partir d'un template
        """
        print(f"\n{'='*50}")
        print(f"🔍 DÉBUT GÉNÉRATION DANS LE SERVICE")
        print(f"{'='*50}")
        
        # Vérifications du template
        template_path = modele.fichier_template.path
        print(f"📂 Template path: {template_path}")
        print(f"  - Existe? {os.path.exists(template_path)}")
        if os.path.exists(template_path):
            print(f"  - Lisible? {os.access(template_path, os.R_OK)}")
            print(f"  - Taille: {os.path.getsize(template_path)} octets")
        
        # Charger le template
        try:
            print("🔄 Chargement du document Word...")
            doc = Document(template_path)
            print(f"✅ Document chargé: {len(doc.paragraphs)} paragraphes")
        except Exception as e:
            print(f"❌ ERREUR chargement: {str(e)}")
            raise
        
        # Collecter toutes les données
        print("🔄 Collecte des données...")
        donnees = {}
        donnees.update(self.collecter_donnees_entreprise(entreprise))
        donnees.update(self.collecter_donnees_opportunite(opportunite, opportunite_type))
        
        # Données système
        donnees.update({
            '{{DATE_JOUR}}': datetime.now().strftime('%d/%m/%Y'),
            '{{HEURE}}': datetime.now().strftime('%H:%M'),
        })
        
        # Données supplémentaires
        if donnees_supp:
            print(f"📝 Données supplémentaires: {donnees_supp}")
            for key, value in donnees_supp.items():
                donnees[f'{{{{{key}}}}}'] = value
        
        print(f"📊 Total variables: {len(donnees)}")
        
        # Remplacer dans tout le document
        print("🔄 Remplacement des variables...")
        self._remplacer_dans_document(doc, donnees)
        print("✅ Remplacement terminé")
        
        # Générer nom fichier unique
        hash_obj = hashlib.md5(f"{entreprise.id}_{opportunite.id}_{datetime.now()}".encode())
        nom_fichier = f"{hash_obj.hexdigest()[:10]}_{modele.nom[:30]}.docx"
        nom_fichier = nom_fichier.replace(' ', '_').replace('/', '_')
        print(f"📄 Nom fichier: {nom_fichier}")
        
        chemin = os.path.join(self.output_dir, nom_fichier)
        print(f"💾 Sauvegarde dans: {chemin}")
        
        # Sauvegarder
        doc.save(chemin)
        
        # Vérifier la sauvegarde
        if os.path.exists(chemin):
            taille = os.path.getsize(chemin)
            print(f"✅ Fichier sauvegardé: {taille} octets")
        else:
            print(f"❌ ERREUR: Fichier non créé!")
        
        print(f"{'='*50}\n")
        
        return chemin, nom_fichier, taille

    def _remplacer_dans_document(self, doc, donnees):
        """Remplace les variables dans le document"""
        # Paragraphes
        for paragraph in doc.paragraphs:
            self._remplacer_paragraphe(paragraph, donnees)
        
        # Tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._remplacer_paragraphe(paragraph, donnees)
        
        # En-têtes et pieds de page
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._remplacer_paragraphe(paragraph, donnees)
            for paragraph in section.footer.paragraphs:
                self._remplacer_paragraphe(paragraph, donnees)
    
    def _remplacer_paragraphe(self, paragraph, donnees):
        """Remplace dans un paragraphe en conservant le formatage"""
        for placeholder, valeur in donnees.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(valeur))